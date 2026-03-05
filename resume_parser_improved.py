#!/usr/bin/env python3
"""
Enhanced Resume Parser
A comprehensive resume parsing solution with improved accuracy and API capabilities.

File structure (all in one file):
  1. Imports & setup (logging, spaCy)
  2. Constants — regexes, section keywords, job/skill/employment/department keywords
  3. Data models — ContactInfo, JobExperience, Education, ParsedResume
  4. LocationValidator — location validation (US states, countries)
  5. ResumeParser — text extraction, contact, sections, skills, experience, education,
     projects, confidence, parse_resume(), to_dict()
"""

# Standard library
import os
import re
import logging
from dataclasses import asdict, dataclass
from datetime import datetime
from typing import Any, Dict, List, Optional

# Third-party
import docx
import pdfplumber
import spacy
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
import nltk
from nltk.corpus import stopwords

# -----------------------------------------------------------------------------
# Logging & NLP setup
# -----------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# NLTK data is downloaded during Docker build
# No need to download at runtime in Lambda

# Load spaCy model (optional - parser will work without it)
nlp = None
try:
    nlp = spacy.load("en_core_web_sm")
    logger.info("SpaCy model 'en_core_web_sm' loaded successfully")
except OSError:
    logger.warning("SpaCy model 'en_core_web_sm' not found. Parser will use pattern-based extraction only. To enable NLP features, run: python -m spacy download en_core_web_sm")
except Exception as e:
    logger.warning(f"Failed to load spaCy model: {e}. Parser will use pattern-based extraction only.")

# =============================================================================
# CONSTANTS — Contact & URL patterns
# =============================================================================
PHONE_REGEX = re.compile(
    r'(?:\+?1[-.\s]?)?(?:\(?([0-9]{3})\)?[-.\s]?)?([0-9]{3})[-.\s]?([0-9]{4})'
    r'|(?:\+?[1-9]\d{0,3}[-.\s]?)?(?:\(?([0-9]{2,4})\)?[-.\s]?)?([0-9]{2,4})[-.\s]?([0-9]{2,4})[-.\s]?([0-9]{2,4})'
    r'|(?:\+?[1-9]\d{0,3}[-.\s]?)?([0-9]{7,15})'
)

EMAIL_REGEX = re.compile(
    r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
)

# LinkedIn profile URLs (with or without protocol/www)
LINKEDIN_REGEX = re.compile(
    r'(?:https?://)?(?:www\.)?linkedin\.com/[^\s,;]+',
    re.IGNORECASE
)

# GitHub profile URLs
GITHUB_REGEX = re.compile(
    r'(?:https?://)?(?:www\.)?github\.com/[A-Za-z0-9_-]+/?',
    re.IGNORECASE
)

# Facebook profile/page URLs
FACEBOOK_REGEX = re.compile(
    r'(?:https?://)?(?:www\.)?(?:fb\.me|facebook\.com|fb\.com)/[^\s,;)+]+',
    re.IGNORECASE
)

# WhatsApp: wa.me links or "WhatsApp" / "WhatsApp:" followed by phone
WHATSAPP_REGEX = re.compile(
    r'(?:https?://)?(?:wa\.me|api\.whatsapp\.com/send\?phone=)/?(\d{7,15})',
    re.IGNORECASE
)

# -----------------------------------------------------------------------------
# Date patterns (used for experience/education date ranges)
# -----------------------------------------------------------------------------
# Apostrophe-like chars (ASCII + Unicode curly/smart quotes from Word/DOCX)
APOS = r"[\'\u2018\u2019]"

# Enhanced date patterns
DATE_PATTERNS = [
    r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*' + APOS + r'?\s*\d{4}',  # Jan 2020, Oct'2020 (with space)
    r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\d{4}',  # Jan2020, July2022 (no space)
    r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*' + APOS + r'?\s*\d{2}(?!\d)',  # Jan'22, May'16 (2-digit year, Word curly apostrophe)
    r'(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}',  # Full month names
    r'\d{1,2}[/-]\d{4}',  # 05/2022, 08-2021 (month/year format)
    r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}',  # 01/15/2020, 12-31-2020
    r'\d{4}[/-]\d{1,2}[/-]\d{1,2}',  # 2020/01/15, 2020-12-31
    r'\d{4}',
    r'(?:present|current|now|till\s+date|till\s+now|tilldate|tillnow)'
]

DATE_RANGE_REGEX = re.compile(
    r'(?i)(?:' + '|'.join(DATE_PATTERNS) + r')\s*(?:-|\u2013|\u2014|\s*to\s*|\s*until\s*|\s*till\s*)\s*(?:' + '|'.join(DATE_PATTERNS) + r'|present|current|now|till\s+date|till\s+now|tilldate|tillnow)',
    re.IGNORECASE
)

# -----------------------------------------------------------------------------
# Section detection keywords
# -----------------------------------------------------------------------------
SECTION_KEYWORDS = {
    "contact": [
        "contact", "contact information", "personal information", "personal details"
    ],
    "professionalSummary": [
        "professional summary", "summary", "profile", "career summary",
        "professional profile", "executive summary", "objective", "career objective"
    ],
    "experience": [
        "projects/technologies experience", "technologies experience", "projects experience",
        "experience", "professional experience", "work experience",
        "employment history", "career history", "work history", "employment",
        "professional background", "work background"
    ],
    "education": [
        "education", "academic background", "academic qualifications",
        "academic achievements", "educational background", "academic credentials",
        "degrees", "qualifications"
    ],
    "skills": [
        "skills", "technical skills", "core competencies", "competencies",
        "expertise", "technologies", "tools", "programming languages"
    ],
    "certifications": [
        "certifications", "certification", "licenses", "license",
        "credentials", "professional certifications"
    ],
    "projects": [
        "projects", "key projects", "notable projects", "project experience",
        "recent projects", "personal projects", "academic projects", "project details"
    ]
}

# Subheadings inside experience (e.g. "Technologies", "Tools") — do not treat as section headers.
EXPERIENCE_SUBSECTION_HEADERS = frozenset([
    "technologies", "tools", "environment", "key technologies", "technical environment",
    "operating systems", "development practices", "key skills", "technologies used",
    "core technologies", "tech stack",     "methodologies"
])

# -----------------------------------------------------------------------------
# Job title & technical skills keywords (O(1) lookups)
# -----------------------------------------------------------------------------
JOB_TITLE_KEYWORDS = {
    "engineer", "developer", "manager", "consultant", "analyst", "specialist",
    "architect", "lead", "officer", "director", "scientist", "designer",
    "administrator", "associate", "intern", "principal", "senior", "junior",
    "staff", "product", "program", "coordinator", "supervisor", "executive",
    "vice president", "ceo", "cto", "cfo", "founder", "co-founder"
}

# Comprehensive skills database (set for O(1) lookups)
TECH_SKILLS = {
    # Programming Languages
    "python", "java", "javascript", "typescript", "c++", "c#", "php", "ruby", "go", "rust", "swift", "kotlin", "scala", "r", "matlab", "perl", "bash", "powershell",
    # Web Technologies
    "react", "angular", "vue", "node.js", "express", "django", "flask", "spring", "laravel", "symfony", "asp.net", "jquery", "bootstrap", "sass", "less", "webpack", "babel",
    # Cloud & DevOps
    "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "terraform", "ansible", "chef", "puppet", "git", "github", "gitlab", "bitbucket", "ci/cd", "devops", "microservices",
    # Databases
    "sql", "mysql", "postgresql", "mongodb", "redis", "cassandra", "elasticsearch", "oracle", "sqlite", "mariadb", "dynamodb", "neo4j", "influxdb",
    # Testing & QA
    "selenium", "cypress", "jest", "mocha", "junit", "testng", "pytest", "postman", "jmeter", "loadrunner", "katalon", "appium", "testing", "qa", "automation", "tosca",
    # Salesforce & CRM
    "salesforce", "apex", "visualforce", "lightning", "salesforce admin", "salesforce developer", "salesforce architect", "crm", "marketo", "hubspot",
    # Mobile Development
    "android", "ios", "react native", "flutter", "xamarin", "ionic", "cordova", "phonegap",
    # Data Science & Analytics
    "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "keras", "spark", "hadoop", "hive", "pig", "kafka", "airflow", "jupyter", "tableau", "power bi", "excel",
    # Operating Systems
    "linux", "windows", "macos", "unix", "ubuntu", "centos", "redhat", "debian", "fedora",
    # Methodologies & Frameworks
    "agile", "scrum", "kanban", "waterfall", "lean", "six sigma", "itil", "pmp", "prince2",
    # Other Technologies
    "api", "rest", "soap", "graphql", "json", "xml", "yaml", "html", "css", "nginx", "apache", "tomcat", "iis"
}

# =============================================================================
# DATA MODELS
# =============================================================================
@dataclass
class ContactInfo:
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    linkedin: Optional[str] = None
    github: Optional[str] = None
    whatsapp: Optional[str] = None
    date_of_birth: Optional[str] = None
    facebook: Optional[str] = None
    gender: Optional[str] = None


# Employment type keywords (order matters for matching; first match wins)
EMPLOYMENT_TYPE_PATTERNS = [
    ("Full-time", re.compile(r"\bfull[- ]?time\b", re.IGNORECASE)),
    ("Part-time", re.compile(r"\bpart[- ]?time\b", re.IGNORECASE)),
    ("Contract", re.compile(r"\bcontract(or)?\b", re.IGNORECASE)),
    ("Internship", re.compile(r"\bintern(ship)?\b", re.IGNORECASE)),
    ("Freelance", re.compile(r"\bfreelance\b", re.IGNORECASE)),
    ("Remote", re.compile(r"\bremote\b", re.IGNORECASE)),
    ("Hybrid", re.compile(r"\bhybrid\b", re.IGNORECASE)),
    ("Volunteer", re.compile(r"\bvolunteer\b", re.IGNORECASE)),
]

# Common department names to detect from context
DEPARTMENT_KEYWORDS = [
    "engineering", "sales", "marketing", "product", "operations", "hr", "human resources",
    "finance", "design", "research", "r&d", "rd", "it", "support", "legal",     "quality",
]

# -----------------------------------------------------------------------------
# Experience, Education, ParsedResume
# -----------------------------------------------------------------------------
@dataclass
class JobExperience:
    title: Optional[str] = None
    company: Optional[str] = None
    location: Optional[str] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    duration: Optional[str] = None
    responsibilities: List[str] = None
    employment_type: Optional[str] = None  # e.g. Full-time, Part-time, Contract
    department: Optional[str] = None  # e.g. Engineering, Sales

    def __post_init__(self):
        if self.responsibilities is None:
            self.responsibilities = []

@dataclass
class Education:
    degree: Optional[str] = None
    field: Optional[str] = None
    institution: Optional[str] = None
    location: Optional[str] = None
    graduation_date: Optional[str] = None
    gpa: Optional[str] = None

@dataclass
class ParsedResume:
    file_path: str
    contact: ContactInfo
    professionalSummary: Optional[str] = None
    experience: List[JobExperience] = None
    education: List[Education] = None
    skills: List[str] = None
    certifications: List[str] = None
    projects: List[str] = None
    raw_text: Optional[str] = None
    confidence_score: float = 0.0

    def __post_init__(self):
        if self.experience is None:
            self.experience = []
        if self.education is None:
            self.education = []
        if self.skills is None:
            self.skills = []
        if self.certifications is None:
            self.certifications = []
        if self.projects is None:
            self.projects = []


# =============================================================================
# LOCATION VALIDATOR
# =============================================================================
class LocationValidator:
    """Location validator using location database for accurate validation."""
    
    # US State abbreviations
    US_STATES = {
        'AL', 'AK', 'AZ', 'AR', 'CA', 'CO', 'CT', 'DE', 'FL', 'GA',
        'HI', 'ID', 'IL', 'IN', 'IA', 'KS', 'KY', 'LA', 'ME', 'MD',
        'MA', 'MI', 'MN', 'MS', 'MO', 'MT', 'NE', 'NV', 'NH', 'NJ',
        'NM', 'NY', 'NC', 'ND', 'OH', 'OK', 'OR', 'PA', 'RI', 'SC',
        'SD', 'TN', 'TX', 'UT', 'VT', 'VA', 'WA', 'WV', 'WI', 'WY', 'DC'
    }
    
    # US State full names
    US_STATE_NAMES = {
        'alabama', 'alaska', 'arizona', 'arkansas', 'california', 'colorado',
        'connecticut', 'delaware', 'florida', 'georgia', 'hawaii', 'idaho',
        'illinois', 'indiana', 'iowa', 'kansas', 'kentucky', 'louisiana',
        'maine', 'maryland', 'massachusetts', 'michigan', 'minnesota',
        'mississippi', 'missouri', 'montana', 'nebraska', 'nevada',
        'new hampshire', 'new jersey', 'new mexico', 'new york',
        'north carolina', 'north dakota', 'ohio', 'oklahoma', 'oregon',
        'pennsylvania', 'rhode island', 'south carolina', 'south dakota',
        'tennessee', 'texas', 'utah', 'vermont', 'virginia', 'washington',
        'west virginia', 'wisconsin', 'wyoming', 'district of columbia'
    }
    
    # Common country names (most frequently appearing in resumes)
    COMMON_COUNTRIES = {
        'usa', 'united states', 'us', 'canada', 'uk', 'united kingdom',
        'australia', 'india', 'germany', 'france', 'spain', 'italy',
        'netherlands', 'sweden', 'norway', 'denmark', 'switzerland',
        'singapore', 'japan', 'china', 'south korea', 'brazil', 'mexico'
    }
    
    def __init__(self):
        """Initialize location validator."""
        self._country_cache = {}
        # Try to load pycountry if available (optional dependency)
        try:
            import pycountry
            self.pycountry = pycountry
            self.has_pycountry = True
        except ImportError:
            self.pycountry = None
            self.has_pycountry = False
            logger.info("pycountry not available, using basic country validation")
    
    def is_valid_us_state(self, state: str) -> bool:
        """Check if string is a valid US state abbreviation or name."""
        state_upper = state.upper().strip()
        state_lower = state.lower().strip()
        return state_upper in self.US_STATES or state_lower in self.US_STATE_NAMES
    
    def is_valid_country(self, country: str) -> bool:
        """Check if string is a valid country name."""
        country_lower = country.lower().strip()
        
        # Check common countries first (fast path)
        if country_lower in self.COMMON_COUNTRIES:
            return True
        
        # Use pycountry if available
        if self.has_pycountry:
            if country_lower in self._country_cache:
                return self._country_cache[country_lower]
            
            try:
                # Check by name
                result = self.pycountry.countries.search_fuzzy(country)
                if result:
                    self._country_cache[country_lower] = True
                    return True
            except LookupError:
                pass
            
            # Check by alpha_2 code
            try:
                if self.pycountry.countries.get(alpha_2=country.upper()):
                    self._country_cache[country_lower] = True
                    return True
            except (KeyError, AttributeError):
                pass
            
            self._country_cache[country_lower] = False
            return False
        
        # Fallback: check if it looks like a country name (proper noun, reasonable length)
        if re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*$', country) and 3 <= len(country) <= 30:
            return True
        
        return False
    
    def validate_location(self, candidate: str) -> bool:
        """Validate a location candidate against location database."""
        if not candidate or len(candidate) < 3 or len(candidate) > 100:
            return False
        
        candidate = candidate.strip()
        
        # Pattern: City, State (e.g., "Atlanta, GA")
        city_state_pattern = r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),\s*([A-Z]{2})$'
        match = re.match(city_state_pattern, candidate)
        if match:
            city, state = match.groups()
            # Validate state
            if self.is_valid_us_state(state):
                return True
        
        # Pattern: City, State ZIP (e.g., "Atlanta, GA 30309")
        city_state_zip_pattern = r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),\s*([A-Z]{2})\s+\d{5}(?:-\d{4})?$'
        match = re.match(city_state_zip_pattern, candidate)
        if match:
            city, state = match.groups()
            if self.is_valid_us_state(state):
                return True
        
        # Pattern: City, Country/State (e.g., "London, UK" or "Toronto, Canada")
        if ',' in candidate:
            parts = [p.strip() for p in candidate.split(',')]
            if len(parts) == 2:
                city_part, region_part = parts
                
                # City part must be proper noun format
                if not re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*$', city_part):
                    return False
                
                # Check if region is US state
                if self.is_valid_us_state(region_part):
                    return True
                
                # Check if region is a country
                if self.is_valid_country(region_part):
                    return True
        
        return False


# =============================================================================
# RESUME PARSER
# =============================================================================
class ResumeParser:
    """Enhanced resume parser with improved accuracy and modularity."""

    # -------------------------------------------------------------------------
    # Initialization
    # -------------------------------------------------------------------------
    def __init__(self, config: Optional[Dict] = None):
        self.config = config or {}
        self.nlp = nlp  # Store spaCy model for use in class methods
        self.location_validator = LocationValidator()  # Initialize location validator
        
        # Pre-compile frequently used regex patterns for performance
        self._compiled_patterns = {
            'role_prefix': re.compile(r'^(?:Role|Position)[:\-]\s*', re.IGNORECASE),
            'role_in_text': re.compile(r'(?:Role|Position)[:\-]\s*(.+?)$', re.IGNORECASE),
            'client_prefix': re.compile(r'^Client[:\-]\s*', re.IGNORECASE),
            'client_pattern': re.compile(r'Client[:\-]\s*(.+?)(?:\s*[,\-]|$)', re.IGNORECASE),
            'location_pattern': re.compile(r',\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*,\s*[A-Z]{2}|[A-Z]{2})\s*$'),
            'bullet_point': re.compile(r'^[●•\-\*]'),
            'responsibility_start': re.compile(r'^(?:Responsible|Worked|Developed|Created|Implemented|Environment)[:\-]', re.IGNORECASE),
            'description_start': re.compile(r'^(?:Description|Responsibilities?|Environment)[:\-]', re.IGNORECASE),
            'company_indicators': re.compile(r'\b(?:Inc|LLC|Corp|Co|Ltd|Company|Corporation|Bank|Services|Solutions|Technologies|Systems)\b', re.IGNORECASE),
            'proper_noun': re.compile(r'^[A-Z][A-Za-z0-9\s&.,\-]+$'),
            'year_pattern': re.compile(r'\b(19|20)\d{2}\b'),
            'gpa_patterns': [
                re.compile(r'(?:GPA|CGPA)[:\-]?\s*(\d+\.?\d*)', re.IGNORECASE),
                re.compile(r'(?:GPA|CGPA)[:\-]?\s*(\d+\.?\d*)\s*/\s*\d+\.?\d*', re.IGNORECASE),
                re.compile(r'(?:GPA|CGPA)[:\-]?\s*(\d+\.?\d*)\s*(?:out\s+of|/)\s*\d+\.?\d*', re.IGNORECASE),
                re.compile(r'(\d+\.?\d*)\s*/\s*\d+\.?\d*'),
                re.compile(r'(\d+\.?\d*)\s*(?:GPA|CGPA)', re.IGNORECASE),
                re.compile(r'(\d+\.?\d*)\s*(?:GPA|CGPA)\s*\([^)]*\)', re.IGNORECASE),
            ],
        }
        
        # Cache job title keywords as set for fast lookups
        self._job_title_keywords_set = JOB_TITLE_KEYWORDS
        
        # Cache tech skills as set for fast lookups
        self._tech_skills_set = TECH_SKILLS
        
        try:
            self.stop_words = set(stopwords.words('english'))
        except LookupError:
            # Fallback to basic English stopwords if NLTK data is not available
            logger.warning("NLTK stopwords not found, using fallback stopwords")
            self.stop_words = {
                'a', 'an', 'and', 'are', 'as', 'at', 'be', 'by', 'for', 'from',
                'has', 'he', 'in', 'is', 'it', 'its', 'of', 'on', 'that', 'the',
                'to', 'was', 'will', 'with', 'i', 'you', 'we', 'they', 'this',
                'these', 'those', 'or', 'but', 'if', 'when', 'where', 'why',
                'how', 'all', 'any', 'both', 'each', 'few', 'more', 'most',
                'other', 'some', 'such', 'no', 'nor', 'not', 'only', 'own',
                'same', 'so', 'than', 'too', 'very', 'can', 'could', 'should',
                'would', 'may', 'might', 'must', 'shall', 'do', 'does', 'did',
                'have', 'had', 'having', 'been', 'being', 'am', 'are', 'is',
                'was', 'were', 'be', 'being', 'been', 'have', 'has', 'had',
                'having', 'do', 'does', 'did', 'will', 'would', 'could',
                'should', 'may', 'might', 'must', 'shall'
            }

    # -------------------------------------------------------------------------
    # Helpers — job title, environment, employment type, department, filters
    # -------------------------------------------------------------------------
    def _contains_job_title_keyword(self, text: str) -> bool:
        """Check if text contains any job title keyword (optimized with cached set)."""
        text_lower = text.lower()
        return any(keyword in text_lower for keyword in self._job_title_keywords_set)

    def _looks_like_environment_or_tech_stack(self, text: str) -> bool:
        """Return True if line looks like an Environment/Tools line or comma-separated tech stack, not a company or title."""
        if not text or len(text.strip()) < 15:
            return False
        t = text.strip().lower()
        if re.match(r"^(?:environment|environments?/tools?|tools?\s*used)[\s:]", t):
            return True
        # Comma-separated list of tech terms (JSON, JavaScript, Python, Django, Git, AWS, etc.)
        tech_terms = (
            r"\b(?:json|javascript|jquery|angular|python|django|html|css|bootstrap|git|github|aws|"
            r"mysql|postgres|redis|rabbitmq|jenkins|heroku|azure|linux|restful|api|rest|"
            r"kafka|swagger|rspec|cucumber|mvc|xml|ajax|react|node\.?js)\b"
        )
        parts = [p.strip() for p in re.split(r"[,;]", t) if p.strip()]
        if len(parts) >= 3:
            matches = sum(1 for p in parts if re.search(tech_terms, p, re.IGNORECASE))
            if matches >= 2:
                return True
        return False

    def _extract_employment_type(self, text: str) -> Optional[str]:
        """Extract employment type (Full-time, Part-time, Contract, etc.) from text. Returns first match or None."""
        if not text or not text.strip():
            return None
        for label, pattern in EMPLOYMENT_TYPE_PATTERNS:
            if pattern.search(text):
                return label
        return None

    def _extract_department(self, text: str) -> Optional[str]:
        """Extract department (Engineering, Sales, etc.) from text. Looks for 'Department: X', 'X Department', or known department keywords."""
        if not text or not text.strip():
            return None
        text_lower = text.lower()
        # "Department: Engineering" or "Dept.: Engineering"
        dept_match = re.search(r"department\s*[.:]\s*([a-z][a-z\s&]+?)(?:\s*[,\n]|$)", text_lower, re.IGNORECASE)
        if dept_match:
            return dept_match.group(1).strip().title()
        # "Engineering Department"
        for kw in DEPARTMENT_KEYWORDS:
            if re.search(rf"\b{re.escape(kw)}\s+department\b", text_lower):
                return "R&D" if kw in ("r&d", "rd") else kw.title()
        # Standalone department keyword (word boundary) in short context
        for kw in DEPARTMENT_KEYWORDS:
            if re.search(rf"\b{re.escape(kw)}\b", text_lower):
                if kw in ("r&d", "rd"):
                    return "R&D"
                return kw.title()
        return None

    def _is_date_range_only(self, s: Optional[str]) -> bool:
        """Return True if s looks like only a date range (e.g. '02/07/2023 to 02/06/2025'), not a job title."""
        if not s or not s.strip():
            return False
        t = s.strip()
        # Purely numeric date range: MM/DD/YYYY to MM/DD/YYYY or DD/MM/YYYY to DD/MM/YYYY
        if re.match(r"^\d{1,2}/\d{1,2}/\d{2,4}\s*(?:to|–|-|—)\s*\d{1,2}/\d{1,2}/\d{2,4}\s*$", t, re.IGNORECASE):
            return True
        # Month'YYYY – Month YYYY or Feb'2016 – October 2020
        if re.match(r"^(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*['']?\s*\d{2,4}\s*(?:–|-|—|to)\s*(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)?[a-z]*\s*\d{2,4}\s*$", t, re.IGNORECASE):
            return True
        return False

    def _is_location_or_context_title(self, s: Optional[str]) -> bool:
        """Return True if title looks like a location/context line, e.g. '(OPTUM Technology Location: Eden Prairie MN USA (From April 2019 to October 2020))'."""
        if not s or not s.strip():
            return False
        t = s.strip()
        if not (t.startswith("(") and t.endswith(")")):
            return False
        # Parenthetical with "Location:" and "(From ... to ...)"
        if re.search(r"location\s*:\s*.+\(from\s+.+\s+to\s+.+\)", t, re.IGNORECASE):
            return True
        # (Company Name: (From Date to Date)) or (OPTUM Technology India: (From Feb 2006 to April 2019))
        if re.search(r"\(from\s+.+\s+to\s+.+\)\s*\)\s*$", t, re.IGNORECASE):
            return True
        return False

    def _is_certification_or_table_entry(self, job: JobExperience) -> bool:
        """Return True if this job entry is likely a certification row or table header, not real work experience."""
        company = (job.company or "").strip().lower()
        title = (job.title or "").strip().lower()
        combined = f"{company} {title}"
        # Certification table header or column names
        if re.search(r"\b(provided\s+by|coverage)\b", combined) and re.search(r"\bcertification\b", combined):
            return True
        if re.search(r"^certification\s+provided\s+by\s+coverage\s*$", company):
            return True
        # Certification name as company (e.g. "SnowPro Core Certification Snowflake", "SnowPro Advanced Architect Snowflake")
        if re.search(r"\bsnowpro\s+(?:core|advanced|architect)\b", combined):
            return True
        if re.search(r"\b(?:certified|certification)\s+(?:in|by|provided)\b", combined):
            return True
        # Company is just a product/vendor name often used in cert tables (e.g. "Snowflake" with no real title and cert-like context)
        if company in ("snowflake", "salesforce") and self._is_date_range_only(job.title):
            return True
        # Title is only a date range and company looks like cert/table (short or contains Certification)
        if self._is_date_range_only(job.title):
            if re.search(r"\bcertification\b", company) or len(company.split()) <= 2:
                return True
        # Company is a sentence fragment (e.g. "Client X and Y, Since September 2022.")
        if re.search(r",\s*since\s+\w+\s+\d{4}\s*\.?\s*$", company):
            return True
        # Title is only a certification phrase (e.g. "Administrator Certified.") and company looks like product/vendor
        if title and len(title.split()) <= 4 and re.search(r"(?:certified|certification)\.?\s*$", title):
            if re.search(r"\b(salesforce|snowflake|platform\s+system)\b", company):
                return True
        # Empty company and title is only a date range or a location/context line (e.g. "(OPTUM Technology Location: ... (From ... to ...))")
        if not company.strip():
            if self._is_date_range_only(job.title) or self._is_location_or_context_title(job.title):
                return True
        return False

    # -------------------------------------------------------------------------
    # Text extraction (PDF, DOCX, table formatting)
    # -------------------------------------------------------------------------
    def extract_text_from_pdf(self, path: str) -> str:
        """Extract text from PDF with improved error handling.
        
        Extracts both regular text and table content to ensure tabular data
        (common in skills, education, work history sections) is captured.
        """
        try:
            text_parts = []
            with pdfplumber.open(path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        # Extract regular text
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                        
                        # Extract tables - common in resumes for skills, education, etc.
                        try:
                            tables = page.extract_tables()
                            if tables:
                                for table in tables:
                                    if table:
                                        table_text = self._format_table_as_text(table)
                                        if table_text:
                                            text_parts.append(table_text)
                        except Exception as e:
                            logger.warning(f"Error extracting tables from page {page_num + 1}: {e}")
                            
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                        continue
            return "\n".join(text_parts).strip()
        except Exception as e:
            logger.error(f"Error reading PDF {path}: {e}")
            return ""
    
    def _format_table_as_text(self, table: List[List[str]]) -> str:
        """Convert a table (list of rows) into readable text format.
        
        Handles common resume table formats:
        - Two-column tables (label: value)
        - Multi-column skill/technology tables
        - Education/experience tables with dates
        """
        if not table:
            return ""
        
        lines = []
        for row in table:
            if not row:
                continue
            # Filter out None/empty cells and strip whitespace
            cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if cells:
                # Join cells with appropriate separator
                # For two-column tables (common label:value format), use colon
                if len(cells) == 2:
                    lines.append(f"{cells[0]}: {cells[1]}")
                else:
                    # For multi-column, use pipe or comma separation
                    lines.append(" | ".join(cells))
        
        return "\n".join(lines)

    def _normalize_raw_text(self, text: str) -> str:
        """Normalize raw extracted text for consistent parsing: Unicode spaces, newlines, BOM."""
        if not text or not text.strip():
            return text or ""
        # Replace non-breaking and other common Unicode spaces with ASCII space
        text = re.sub(r"[\u00a0\u2000-\u200b\u202f\u205f\ufeff]+", " ", text)
        # Collapse multiple newlines to at most two (preserve paragraph breaks)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def extract_text_from_docx(self, path: str) -> str:
        """Extract text from DOCX with improved error handling.

        Note: Many modern resume templates put the candidate's name, email, and
        phone number in the document header. By default, `python-docx` exposes
        body paragraphs via `doc.paragraphs` but header/footer text must be
        read explicitly from `section.header` / `section.footer`.
        
        Also extracts table content which is common in resumes for skills,
        education, and work history sections.
        """
        try:
            doc = docx.Document(path)
            parts: List[str] = []

            # 1) Collect header text (where contact info often lives)
            # We collect from all sections, de-duplicating later.
            try:
                for section in doc.sections:
                    header = section.header
                    if not header:
                        continue
                    for para in header.paragraphs:
                        text = para.text.strip()
                        if text:
                            parts.append(text)
            except Exception as e:
                # Header parsing issues shouldn't break entire extraction
                logger.warning(f"Error reading DOCX headers for {path}: {e}")

            # 2) Collect body content in document order (paragraphs and tables interleaved).
            # Using doc.paragraphs + doc.tables separately loses order: all paragraphs come
            # first, then all tables, so section detection (find_sections) misassigns content.
            # Iterating body elements in order makes DOCX behave like PDF and fixes missing data.
            try:
                for child in doc.element.body:
                    if isinstance(child, CT_P):
                        para = DocxParagraph(child, doc)
                        text = para.text.strip()
                        if text:
                            parts.append(text)
                    elif isinstance(child, CT_Tbl):
                        table = DocxTable(child, doc)
                        table_data = []
                        for row in table.rows:
                            row_data = []
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                if cell_text:
                                    row_data.append(cell_text)
                            if row_data:
                                table_data.append(row_data)
                        if table_data:
                            table_text = self._format_table_as_text(table_data)
                            if table_text:
                                parts.append(table_text)
            except Exception as e:
                logger.warning(f"Error reading DOCX body in order for {path}: {e}")

            # 4) Optionally collect footer text (rarely contains contact info,
            # but cheap to include and may add missing details)
            try:
                for section in doc.sections:
                    footer = section.footer
                    if not footer:
                        continue
                    for para in footer.paragraphs:
                        text = para.text.strip()
                        if text:
                            parts.append(text)
            except Exception as e:
                logger.warning(f"Error reading DOCX footers for {path}: {e}")

            # De-duplicate while preserving order to avoid repeated header text
            seen = set()
            unique_parts: List[str] = []
            for t in parts:
                if t not in seen:
                    seen.add(t)
                    unique_parts.append(t)

            return "\n".join(unique_parts).strip()
        except Exception as e:
            logger.error(f"Error reading DOCX {path}: {e}")
            return ""

    # -------------------------------------------------------------------------
    # Contact — name, email verification, contact info, location
    # -------------------------------------------------------------------------
    def extract_name(self, text: str) -> str | None:
        """Extract name from resume text with improved accuracy using pattern-based validation."""
        
        def is_valid_name(candidate: str) -> bool:
            """Validate if a candidate string is likely a real name using pattern-based checks."""
            if not candidate:
                return False
            
            candidate = candidate.strip()
            words = candidate.split()
            
            # Structural validation: name length (2-4 words typical for names)
            if not (2 <= len(words) <= 4):
                return False
            
            # Pattern validation: names should be primarily alphabetic (allow hyphens and apostrophes)
            cleaned = re.sub(r"['-]", "", candidate)
            if not cleaned.replace(" ", "").isalpha():
                return False
            
            # Pattern validation: proper noun format - each word should start with capital
            original_words = candidate.split()
            for word in original_words:
                if not word[0].isupper():
                    return False
                # Names don't have mixed case within words (e.g., "JavaScript" is not a name)
                if len(word) > 1 and re.search(r'[a-z][A-Z]', word):
                    return False
            
            # Pattern validation: reject if it contains technical patterns
            # Check for common technical suffixes/patterns
            candidate_lower = candidate.lower()

            # Reject obvious job-title phrases like "Front End UI Developer", "Software Engineer", etc.
            # We keep this pattern small and focused so it stays maintainable.
            if re.search(r"\b(developer|engineer|manager|consultant|analyst|tester|architect|lead|designer|specialist|intern)\b", candidate_lower):
                return False
            
            # Reject words ending in "ing" (gerunds/verbs like "Testing", "Loading", "Processing")
            # These are almost never names
            for word in words:
                word_lower = word.lower()
                if word_lower.endswith('ing') and len(word_lower) >= 5:
                    return False
            
            # Reject if it contains other technical word patterns
            # Technical terms often have specific patterns (e.g., "Framework", "API")
            # Check for words ending in common technical suffixes
            technical_suffixes = ['tion', 'ment', 'ness', 'ity', 'ism']
            for word in words:
                word_lower = word.lower()
                # Very long words (>12 chars) with technical suffixes are suspicious
                if len(word_lower) > 12:
                    for suffix in technical_suffixes:
                        if word_lower.endswith(suffix) and len(word_lower) > len(suffix) + 5:
                            return False
            
            # Pattern validation: reject concatenated words (no spaces in long strings)
            # Names should have proper spacing between words
            for word in words:
                # Very long words (>15 chars) are suspicious for names
                if len(word) > 15:
                    return False
                # Check for mixed case patterns that suggest concatenation
                if re.search(r'[A-Z][a-z]{5,}[A-Z]', word):
                    return False
            
            # Pattern validation: reject if it looks like a section header
            # Section headers often have all caps or specific patterns, but many resumes
            # use ALL CAPS for the candidate name on the first line (e.g., "BALRAJ REDDY").
            if candidate.isupper():
                # Allow short 2–3 word ALL CAPS strings that look like real names
                header_like_keywords = {
                    "professional", "summary", "experience", "education",
                    "skills", "projects", "certifications", "objective",
                    "profile", "resume", "curriculum", "vitae"
                }
                word_set = {w for w in words}
                # If it obviously contains header keywords, reject
                if word_set & {w.upper() for w in header_like_keywords}:
                    return False
                # Otherwise, only reject long ALL CAPS phrases; keep short ones as valid
                if len(words) > 3 or len(candidate) > 40:
                    return False
            
            # Pattern validation: reject if it contains numbers or special chars (except hyphens/apostrophes)
            if re.search(r'[0-9]', candidate):
                return False
            
            # Pattern validation: reject if words are too short (single letter words are rare in names)
            for word in words:
                if len(word) == 1 and word.lower() not in ['i', 'o']:  # Allow 'I' and 'O' as initials
                    return False
            
            return True
        
        # Only scan top section (first 8 lines)
        header = "\n".join(text.split("\n")[:8]).strip()
        lines = [line.strip() for line in header.split("\n") if line.strip()]
        
        # Priority 1: First non-empty line (most resumes have name here)
        if lines:
            first_line = lines[0].strip()
            # Special, slightly looser rule for the very first line: many names include
            # a lowercase middle name or mixed casing (e.g., "Calvin raj Namburi"),
            # and some resumes only have a single given name on the first line (e.g., "Mukesh").
            # Accept if:
            #   - 1–4 alphabetic words
            #   - does NOT contain obvious job-title or section-header keywords
            #   - does not contain digits/symbols
            first_words = first_line.split()
            header_like_keywords = {
                "professional", "summary", "experience", "education",
                "skills", "projects", "certifications", "objective",
                "profile", "resume", "curriculum", "vitae"
            }
            if 2 <= len(first_words) <= 4:
                if all(re.sub(r"['-]", "", w).isalpha() for w in first_words):
                    lower_first = first_line.lower()
                    if not any(kw in lower_first for kw in JOB_TITLE_KEYWORDS) and not any(
                        kw in lower_first for kw in header_like_keywords
                    ):
                        # Looks like a reasonable multi-word name header; accept directly
                        return first_line
            elif len(first_words) == 1:
                single = first_words[0]
                cleaned_single = re.sub(r"['-]", "", single)
                lower_single = cleaned_single.lower()
                # Single-word name heuristic: capitalized, alphabetic, reasonable length, and
                # not an obvious section/header keyword.
                if (
                    cleaned_single.isalpha()
                    and single[0].isupper()
                    and 3 <= len(cleaned_single) <= 20
                    and lower_single not in header_like_keywords
                    and lower_single not in JOB_TITLE_KEYWORDS
                ):
                    return first_line
            # Otherwise fall back to the stricter validator
            if is_valid_name(first_line):
                return first_line
        
        # Priority 2: spaCy NER only on the header (if available)
        if self.nlp:
            try:
                doc = self.nlp(header)
                # spaCy PERSON entities with validation (still constrained to header)
                for ent in doc.ents:
                    if ent.label_ == "PERSON":
                        candidate = ent.text.strip()
                        if is_valid_name(candidate):
                            return candidate
            except Exception as e:
                logger.debug(f"SpaCy NER failed for name extraction: {e}")
        
        # Priority 3: Check remaining lines (2nd–5th) as a fallback
        for line in lines[1:5]:
            if line and is_valid_name(line):
                return line
        
        return None

    def verify_name_with_email(self, name: str, email: str) -> bool:
        """Verify if extracted name matches patterns in email address."""
        if not name or not email or '@' not in email:
            return False
        
        # Get local part of email (before @)
        local_part = email.split('@')[0].lower()
        
        # Normalize name: convert to lowercase and extract name parts
        name_lower = name.lower().strip()
        name_words = [word.lower() for word in name.split() if word.isalpha()]
        
        if len(name_words) < 2:
            return False
        
        # Pattern 1: Check if name words appear in email (with separators)
        # Common separators in emails: ., _, -, (none)
        separators = ['.', '_', '-', '']
        
        for sep in separators:
            if sep:
                # Check if email contains name parts separated by delimiter
                # e.g., "john.doe" matches "John Doe"
                email_parts = local_part.split(sep)
                if len(email_parts) >= 2:
                    # Check if at least 2 name words appear in email parts
                    matches = 0
                    for name_word in name_words:
                        for email_part in email_parts:
                            # Exact match or email part starts with name word (e.g., "john" matches "johnsmith")
                            if name_word == email_part or email_part.startswith(name_word):
                                matches += 1
                                break
                    if matches >= 2:
                        return True
        
        # Pattern 2: Check if name words appear consecutively in email (no separator)
        # e.g., "johndoe" matches "John Doe"
        name_concat = ''.join(name_words)
        if name_concat in local_part:
            return True
        
        # Pattern 3: Check if first letters of name words match email pattern
        # e.g., "jd" or "jdoe" matches "John Doe"
        first_letters = ''.join([word[0] for word in name_words])
        if first_letters in local_part and len(first_letters) >= 2:
            return True
        
        # Pattern 4: Check if first name + last name initial or vice versa
        # e.g., "johnd" or "jdoe" matches "John Doe"
        if len(name_words) >= 2:
            first_name = name_words[0]
            last_name = name_words[-1]
            # Pattern: firstname + lastname initial
            if first_name + last_name[0] in local_part:
                return True
            # Pattern: firstname initial + lastname
            if first_name[0] + last_name in local_part:
                return True
        
        # Pattern 5: Check if individual name words appear in email
        # At least 2 name words should appear in email
        matches = 0
        for name_word in name_words:
            if len(name_word) >= 3:  # Only check words with 3+ chars
                if name_word in local_part:
                    matches += 1
        if matches >= 2:
            return True
        
        return False

    def extract_name_from_email(self, email: str) -> str | None:
        """Extract name from email address using pattern-based validation."""
        if not email or '@' not in email:
            return None
        
        local_part = email.split('@')[0].lower()
        
        # Pattern 1: Separated by common delimiters (e.g., "john.doe", "john_doe", "john-doe")
        separators = ['.', '_', '-']
        for sep in separators:
            if sep in local_part:
                parts = local_part.split(sep)
                # Filter out non-name parts
                filtered_parts = []
                for part in parts:
                    # Skip pure numbers, very short parts, or parts with numbers
                    if part.isdigit() or len(part) < 2 or re.search(r'\d', part):
                        continue
                    # Only keep alphabetic parts
                    if part.isalpha():
                        filtered_parts.append(part)
                
                # Need at least 2 parts to form a name
                if len(filtered_parts) >= 2:
                    name_candidate = ' '.join(filtered_parts).title()
                    # Use the same validation as extract_name
                    if self._is_valid_name_format(name_candidate):
                        return name_candidate
        
        # Pattern 2: camelCase (e.g., "JohnDoe" -> "John Doe")
        if re.search(r'[a-z][A-Z]', local_part):
            parts = re.split(r'([A-Z][a-z]+)', local_part)
            filtered_parts = [p for p in parts if p and len(p) >= 2 and p.isalpha()]
            if len(filtered_parts) >= 2:
                name_candidate = ' '.join(filtered_parts).title()
                if self._is_valid_name_format(name_candidate):
                    return name_candidate
        
        # Pattern 3: All lowercase concatenated (e.g., "bhavaniakuthota" -> "Bhavani Akuthota")
        # Try intelligent splitting for common name patterns
        if local_part.isalpha() and local_part.islower():
            length = len(local_part)
            # Typical names: 8-20 characters when concatenated
            if 8 <= length <= 20:
                # Try splitting at various points
                # Common patterns: first name (4-8 chars) + last name (4-10 chars)
                for split_point in range(4, min(10, length - 3)):
                    first_part = local_part[:split_point]
                    second_part = local_part[split_point:]
                    
                    # Both parts should be reasonable length (3+ chars each)
                    if len(first_part) >= 3 and len(second_part) >= 3:
                        name_candidate = f"{first_part.title()} {second_part.title()}"
                        if self._is_valid_name_format(name_candidate):
                            return name_candidate
        
        return None

    def _is_valid_name_format(self, candidate: str) -> bool:
        """Helper function to validate name format (reused from extract_name logic)."""
        if not candidate:
            return False
        
        candidate = candidate.strip()
        words = candidate.split()
        
        # Structural validation: name length (2-4 words typical for names)
        if not (2 <= len(words) <= 4):
            return False
        
        # Pattern validation: names should be primarily alphabetic
        cleaned = re.sub(r"['-]", "", candidate)
        if not cleaned.replace(" ", "").isalpha():
            return False
        
        # Pattern validation: proper noun format - each word should start with capital
        for word in words:
            if not word[0].isupper():
                return False
            # Names don't have mixed case within words
            if len(word) > 1 and re.search(r'[a-z][A-Z]', word):
                return False
        
        # Pattern validation: reject words ending in "ing" (gerunds/verbs)
        for word in words:
            word_lower = word.lower()
            if word_lower.endswith('ing') and len(word_lower) >= 5:
                return False
        
        # Pattern validation: reject very long words
        for word in words:
            if len(word) > 15:
                return False
            if re.search(r'[A-Z][a-z]{5,}[A-Z]', word):
                return False
        
        # Reject obvious job-title phrases (e.g., "... Developer", "... Engineer") to
        # avoid treating email-based constructs like "Beer Amuideveloper" as names.
        candidate_lower = candidate.lower()
        for kw in JOB_TITLE_KEYWORDS:
            if kw in candidate_lower:
                return False
        
        # Pattern validation: reject if it contains numbers
        if re.search(r'[0-9]', candidate):
            return False
        
        return True

    def extract_contact_info(self, text: str) -> ContactInfo:
        """Extract contact information with improved accuracy."""
        contact = ContactInfo()
        
        # Extract email
        email_match = EMAIL_REGEX.search(text)
        if email_match:
            contact.email = email_match.group(0).lower()
        
        # Extract phone
        phone_matches = PHONE_REGEX.findall(text)
        if phone_matches:
            # Clean and validate phone numbers
            for match in phone_matches:
                phone = ''.join(filter(str.isdigit, ''.join(match)))
                if 7 <= len(phone) <= 15:
                    contact.phone = phone
                    break
        
        # Extract LinkedIn profile (if present anywhere in text)
        linkedin_match = LINKEDIN_REGEX.search(text)
        if linkedin_match:
            linkedin_url = linkedin_match.group(0).strip().rstrip('.,);]')
            # Normalise to include protocol
            if not linkedin_url.lower().startswith(("http://", "https://")):
                linkedin_url = "https://" + linkedin_url
            contact.linkedin = linkedin_url

        # Extract GitHub profile
        github_match = GITHUB_REGEX.search(text)
        if github_match:
            github_url = github_match.group(0).strip().rstrip('.,);]')
            if not github_url.lower().startswith(("http://", "https://")):
                github_url = "https://" + github_url
            contact.github = github_url

        # Extract Facebook profile
        facebook_match = FACEBOOK_REGEX.search(text)
        if facebook_match:
            facebook_url = facebook_match.group(0).strip().rstrip('.,);]')
            if not facebook_url.lower().startswith(("http://", "https://")):
                facebook_url = "https://" + facebook_url
            contact.facebook = facebook_url

        # Extract WhatsApp: wa.me link or "WhatsApp:" / "WhatsApp" followed by phone on same/short line
        wa_match = WHATSAPP_REGEX.search(text)
        if wa_match:
            contact.whatsapp = wa_match.group(0).strip().rstrip('.,);]')
        else:
            # Look for "WhatsApp" or "WhatsApp:" near a phone number (e.g. in first 800 chars)
            head = text[:800] if len(text) > 800 else text
            if re.search(r"\bwhatsapp\b", head, re.IGNORECASE):
                phone_m = PHONE_REGEX.search(head)
                if phone_m:
                    digits = "".join(filter(str.isdigit, phone_m.group(0)))
                    if 7 <= len(digits) <= 15:
                        contact.whatsapp = digits

        # Extract date of birth (DOB / Date of Birth / Born near a date)
        dob_patterns = [
            r"(?:DOB|Date\s+of\s+Birth|Birth\s+Date|Born|D\.O\.B\.?)\s*[:\-]?\s*(\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4})",
            r"(?:DOB|Date\s+of\s+Birth|Birth\s+Date|Born|D\.O\.B\.?)\s*[:\-]?\s*(\d{4}-\d{2}-\d{2})",
            r"(?:DOB|Date\s+of\s+Birth|Birth\s+Date|Born|D\.O\.B\.?)\s*[:\-]?\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s']*\d{1,2},?\s*\d{4})",
            r"(?:DOB|Date\s+of\s+Birth|Birth\s+Date|Born|D\.O\.B\.?)\s*[:\-]?\s*(\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4})",
        ]
        head_contact = text[:1200] if len(text) > 1200 else text  # DOB usually in header
        for pat in dob_patterns:
            m = re.search(pat, head_contact, re.IGNORECASE)
            if m:
                contact.date_of_birth = m.group(1).strip()
                break

        # Extract gender (in contact/header area: "Gender:", "Sex:", or standalone Male/Female)
        gender_patterns = [
            r"(?:Gender|Sex)\s*[:\-]?\s*(Male|Female|M|F|Other|Prefer\s+not\s+to\s+say)",
            r"\b(Male|Female)\b(?=\s*$|\s*\n|\s*[,\|])",
        ]
        for pat in gender_patterns:
            m = re.search(pat, head_contact, re.IGNORECASE)
            if m:
                g = m.group(1).strip()
                if g.upper() in ("M", "F"):
                    contact.gender = "Male" if g.upper() == "M" else "Female"
                else:
                    contact.gender = g.title()
                break
        
        # Extract name using improved method (header-based)
        contact.name = self.extract_name(text)
        
        # # Verify extracted name with email (if both exist)
        # if contact.name and contact.email:
        #     if not self.verify_name_with_email(contact.name, contact.email):
        #         # Name doesn't match email pattern - might be incorrect, try extracting from email
        #         email_name = self.extract_name_from_email(contact.email)
        #         if email_name:
        #             # If email-based name is found and it's different, prefer email-based one
        #             # (email is more reliable for name extraction)
        #             contact.name = email_name
        # elif not contact.name and contact.email:
        # No name found in text, try extracting from email
        if not contact.name and contact.email:
            contact.name = self.extract_name_from_email(contact.email)
        
        # Extract location using improved multi-strategy approach
        contact.location = self.extract_location_enhanced(text, contact.name, contact.email, contact.phone)
        
        return contact
    
    def extract_location_enhanced(self, text: str, name: Optional[str] = None, email: Optional[str] = None, phone: Optional[str] = None) -> Optional[str]:
        """Completely rewritten location extraction for maximum accuracy."""
        
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if not lines:
            return None
        
        # Rejection patterns - things that should NEVER be considered locations
        REJECTION_PATTERNS = [
            # Technical terms
            r'\b(?:JavaScript|TypeScript|Python|Java|C\+\+|API|REST|SOAP|JSON|XML|HTML|CSS|SQL|NoSQL)\b',
            # Job-related terms
            r'\b(?:Analysis|Design|Development|Implementation|Testing|Software|Programming|Engineering)\b',
            # Domain/industry terms
            r'\b(?:Health|IT|Banking|Finance|Retail|E-commerce|Healthcare|Insurance|Telecom)\b',
            # Framework/tool names
            r'\b(?:React|Angular|Vue|Node|Django|Flask|Spring|Bootstrap|jQuery)\b',
            # Common false positives
            r'\b(?:Environment|Tools|Technologies|Skills|Technical|Programming|Languages)\b',
        ]
        
        def should_reject(candidate: str) -> bool:
            """Comprehensive rejection check - reject anything that looks like a false positive."""
            if not candidate or len(candidate) < 3 or len(candidate) > 80:
                return True
            
            candidate_clean = candidate.strip()
            words = candidate_clean.split()
            
            # Reject if matches name
            if name and candidate_clean.lower() == name.lower():
                return True
            
            # Reject if contains numbers (except ZIP codes in proper format)
            if re.search(r'\d', candidate_clean):
                # Only allow if it's a proper ZIP code format
                if not re.search(r'\d{5}(?:-\d{4})?', candidate_clean):
                    return True
            
            # Reject technical patterns (camelCase, ALLCAPS acronyms)
            if re.search(r'[a-z][A-Z]', candidate_clean) or re.search(r'\b[A-Z]{3,}\b', candidate_clean):
                return True
            
            # Reject if words are too long (locations don't have very long words)
            if any(len(word) > 20 for word in words):
                return True
            
            # Reject if matches any rejection pattern
            candidate_lower = candidate_clean.lower()
            for pattern in REJECTION_PATTERNS:
                if re.search(pattern, candidate_clean, re.IGNORECASE):
                    return True
            
            # Reject common job/domain term combinations
            job_terms = {'analysis', 'design', 'development', 'testing', 'software', 'health', 'it', 'banking', 'finance'}
            if len(words) == 2 and all(w.lower() in job_terms for w in words):
                return True
            
            return False
            
        def validate_location_strict(candidate: str) -> bool:
            """Strict validation - must pass all checks."""
            if should_reject(candidate):
                return False
        
            # Must validate against location database
            return self.location_validator.validate_location(candidate)
        
        # PRIORITY 1: Most reliable pattern - "City, State ZIP" (US format)
        # This is the gold standard - if we find this, it's almost certainly correct
        zip_code_pattern = r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),\s*([A-Z]{2})\s+(\d{5}(?:-\d{4})?)\b'
        
        # Search in header area (first 15 lines - most reliable)
        header_text = '\n'.join(lines[:15])
        for match in re.finditer(zip_code_pattern, header_text):
                candidate = match.group(0).strip()
                if validate_location_strict(candidate):
                    return candidate
        
        # PRIORITY 2: "City, State" (US format) - very reliable
        city_state_pattern = r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),\s*([A-Z]{2})\b'
        
        # First check header area
        for match in re.finditer(city_state_pattern, header_text):
            candidate = match.group(0).strip()
            # Additional validation: state must be valid
            state_part = match.group(2)
            if self.location_validator.is_valid_us_state(state_part) and validate_location_strict(candidate):
                return candidate
        
        # PRIORITY 3: Context-aware search near email/phone
        # Locations are often on the same line or adjacent to contact info
        if email or phone:
            search_item = email if email else phone
            search_pos = text.find(search_item)
            
            if search_pos != -1:
                # Get line index
                line_idx = text[:search_pos].count('\n')
                
                # Check same line and 2 lines before/after
                context_start = max(0, line_idx - 2)
                context_end = min(len(lines), line_idx + 3)
                context_lines = lines[context_start:context_end]
                context_text = '\n'.join(context_lines)
                
                # Try ZIP code pattern first (most reliable)
                for match in re.finditer(zip_code_pattern, context_text):
                    candidate = match.group(0).strip()
                    if validate_location_strict(candidate):
                        return candidate
                
                # Then try City, State
                for match in re.finditer(city_state_pattern, context_text):
                    candidate = match.group(0).strip()
                    state_part = match.group(2)
                    if self.location_validator.is_valid_us_state(state_part) and validate_location_strict(candidate):
                        return candidate
        
        # PRIORITY 4: International format "City, Country" or "City, State/Province"
        # Only if we haven't found anything yet
        city_country_pattern = r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*),\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,2})\b'
        
        # Check header area only (international locations are less common)
        for match in re.finditer(city_country_pattern, header_text):
            candidate = match.group(0).strip()
            country_part = match.group(2).strip()
            
            # Validate country/region part
            if (self.location_validator.is_valid_country(country_part) or 
                self.location_validator.is_valid_us_state(country_part)):
                if validate_location_strict(candidate):
                    return candidate
        
        # PRIORITY 5: Extended search (lines 15-25) - only for ZIP code pattern (most reliable)
        if len(lines) > 15:
            extended_text = '\n'.join(lines[15:25])
            for match in re.finditer(zip_code_pattern, extended_text):
                candidate = match.group(0).strip()
                if validate_location_strict(candidate):
                    return candidate
        
        # PRIORITY 6: Optional spaCy NER (if available) - last resort
        # Only use if we haven't found anything with patterns
        if self.nlp:
            try:
                header_doc = self.nlp(header_text)
                location_candidates = []
                
                for ent in header_doc.ents:
                    if ent.label_ in ["GPE", "LOC"]:
                        candidate = ent.text.strip()
                        # Must pass strict validation
                        if validate_location_strict(candidate):
                            location_candidates.append((candidate, ent.start_char))
                
                # Return earliest valid location
                if location_candidates:
                    location_candidates.sort(key=lambda x: x[1])
                    return location_candidates[0][0]
            except Exception as e:
                logger.debug(f"SpaCy NER failed: {e}")
        
        # No valid location found
        return None

    # -------------------------------------------------------------------------
    # Section detection
    # -------------------------------------------------------------------------
    def find_sections(self, text: str) -> Dict[str, str]:
        """Find and extract resume sections with improved accuracy."""
        lines = text.split('\n')
        sections = {}
        current_section = "header"
        sections[current_section] = []
        
        # Create keyword mapping; check longer phrases first so e.g. "projects/technologies experience"
        # is matched as experience, not projects.
        keyword_to_section = {}
        for section, keywords in SECTION_KEYWORDS.items():
            for keyword in keywords:
                keyword_to_section[keyword.lower()] = section
        # Sort by keyword length descending so "projects/technologies experience" wins over "projects"
        keywords_sorted = sorted(keyword_to_section.keys(), key=len, reverse=True)
        
        # Build a sliding window to detect split headers (e.g., "W\nORK\nH\nISTORY")
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if not line_stripped:
                continue
                
            normalized = line_stripped.lower()
            matched_section = None
            
            # Check for section headers in current line (longest match first)
            for keyword in keywords_sorted:
                section_name = keyword_to_section[keyword]
                if (re.search(r'\b' + re.escape(keyword) + r'\b', normalized) and 
                    len(line_stripped) < 50 and 
                    not any(char.isdigit() for char in line_stripped[:10])):
                    matched_section = section_name
                    break
            
            # If no match, check combined with next few lines (for split headers)
            if not matched_section and i < len(lines) - 3:
                combined = " ".join([l.strip() for l in lines[i:i+4] if l.strip()]).lower()
                for keyword in keywords_sorted:
                    section_name = keyword_to_section[keyword]
                    if re.search(r'\b' + re.escape(keyword) + r'\b', combined):
                        matched_section = section_name
                        break
            
            # Don't switch to skills (or away from experience) when the line is a common
            # *subsection* header inside experience (e.g. "Technologies", "Tools", "Environment").
            # Otherwise the experience section is truncated and later jobs are lost.
            if matched_section and current_section == "experience":
                line_lower = normalized.strip()
                # Check if line is exactly or almost exactly a subsection header (with optional colon/dash)
                line_core = re.sub(r'[:\-\s]+$', '', line_lower).strip()
                if line_core in EXPERIENCE_SUBSECTION_HEADERS or any(
                    line_core == h or line_core.startswith(h + " ") or line_core.startswith(h + ":") or line_core.startswith(h + "-")
                    for h in EXPERIENCE_SUBSECTION_HEADERS
                ):
                    matched_section = None
            
            if matched_section:
                current_section = matched_section
                if current_section not in sections:
                    sections[current_section] = []
                continue
            else:
                sections.setdefault(current_section, []).append(line_stripped)
        
        # Clean up sections
        for section in sections:
            content_lines = sections[section]
            # Remove empty lines from start and end
            while content_lines and not content_lines[0].strip():
                content_lines.pop(0)
            while content_lines and not content_lines[-1].strip():
                content_lines.pop()
            sections[section] = "\n".join(content_lines).strip()
        
        return sections

    # -------------------------------------------------------------------------
    # Skills extraction
    # -------------------------------------------------------------------------
    def extract_skills(self, text: str, sections: Optional[Dict[str, str]] = None) -> List[str]:
        """Extract technical skills and competencies with comprehensive detection.

        `sections` can be passed in when the caller has already segmented the
        resume text, to avoid re-running section detection multiple times.
        """
        skills = set()
        text_lower = text.lower()
        # Avoid recomputing sections repeatedly inside this method
        if sections is None:
            sections = self.find_sections(text)
        
        # 1. Extract known technical skills with word boundaries (optimized with cached set)
        for skill in self._tech_skills_set:
            # Use word boundaries to avoid partial matches
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text_lower):
                skills.add(skill.title())
        
        # 2. Extract skills from specific sections
        skills_section = sections.get("skills", "")
        if skills_section:
            # Split by common delimiters
            skill_lines = re.split(r'[,;|\n•\-\*]', skills_section)
            for line in skill_lines:
                line = line.strip()
                if line and len(line) > 2:
                    # Check if it's a known skill (optimized with cached set)
                    line_lower = line.lower()
                    for skill in self._tech_skills_set:
                        if skill.lower() in line_lower:
                            skills.add(skill.title())
                    # Add the line if it looks like a skill
                    if any(keyword in line.lower() for keyword in 
                          ['programming', 'language', 'framework', 'tool', 'technology', 'platform']):
                        skills.add(line.title())
        
        # 3. Extract from experience and summary sections
        experience_text = sections.get("experience", "")
        summary_text = sections.get("professionalSummary", "")
        combined_text = experience_text + " " + summary_text
        
        # Look for skill patterns in experience
        skill_context_patterns = [
            r'(?:proficient|experienced|skilled|expert|knowledgeable|familiar)\s+(?:in|with|at)?\s*([^.,\n]+)',
            r'(?:experience|expertise|skills?)\s+(?:in|with|at)?\s*([^.,\n]+)',
            r'(?:worked\s+with|used|utilized|implemented|developed)\s+([^.,\n]+)',
            r'(?:technologies?|tools?|frameworks?|languages?|platforms?):\s*([^.,\n]+)'
        ]
        
        for pattern in skill_context_patterns:
            matches = re.findall(pattern, combined_text, re.IGNORECASE)
            for match in matches:
                # Split by common delimiters
                potential_skills = re.split(r'[,;|\s+and\s+|\s+&\s+]', match)
                for skill in potential_skills:
                    skill = skill.strip()
                    if skill and len(skill) > 2:
                        # Check if it's a known skill (optimized with cached set)
                        skill_lower = skill.lower()
                        for known_skill in self._tech_skills_set:
                            if known_skill.lower() in skill_lower:
                                skills.add(known_skill.title())
        
        # 4. Extract from bullet points and lists
        bullet_patterns = [
            r'[•\-\*]\s*([^.\n]+(?:python|java|javascript|react|angular|aws|azure|docker|kubernetes|sql|mongodb|salesforce|selenium|testing|automation)[^.\n]*)',
            r'[•\-\*]\s*(?:proficient|experienced|skilled|expert|knowledgeable|familiar)\s+(?:in|with|at)?\s*([^.\n]+)',
            r'[•\-\*]\s*(?:worked\s+with|used|utilized|implemented|developed)\s+([^.\n]+)'
        ]
        
        for pattern in bullet_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                # Split by common delimiters
                potential_skills = re.split(r'[,;|\s+and\s+|\s+&\s+]', match)
                for skill in potential_skills:
                    skill = skill.strip()
                    if skill and len(skill) > 2:
                        # Check if it's a known skill (optimized with cached set)
                        skill_lower = skill.lower()
                        for known_skill in self._tech_skills_set:
                            if known_skill.lower() in skill_lower:
                                skills.add(known_skill.title())
        
        # 5. Extract from specific technology patterns
        tech_patterns = [
            # Programming languages
            r'\b(?:Python|Java|JavaScript|TypeScript|C\+\+|C#|PHP|Ruby|Go|Rust|Swift|Kotlin|Scala|R|MATLAB|Perl|Bash|PowerShell)\b',
            # Web frameworks
            r'\b(?:React|Angular|Vue|Node\.js|Express|Django|Flask|Spring|Laravel|Symfony|ASP\.NET|jQuery|Bootstrap|Sass|Less|Webpack|Babel)\b',
            # Cloud & DevOps
            r'\b(?:AWS|Azure|GCP|Docker|Kubernetes|Jenkins|Terraform|Ansible|Chef|Puppet|Git|GitHub|GitLab|Bitbucket|CI/CD|DevOps|Microservices)\b',
            # Databases
            r'\b(?:SQL|MySQL|PostgreSQL|MongoDB|Redis|Cassandra|Elasticsearch|Oracle|SQLite|MariaDB|DynamoDB|Neo4j|InfluxDB)\b',
            # Testing tools
            r'\b(?:Selenium|Cypress|Jest|Mocha|JUnit|TestNG|PyTest|Postman|JMeter|LoadRunner|Katalon|Appium|TOSCA)\b',
            # Salesforce
            r'\b(?:Salesforce|Apex|Visualforce|Lightning|Salesforce Admin|Salesforce Developer|Salesforce Architect|CRM|Marketo|HubSpot)\b',
            # Mobile
            r'\b(?:Android|iOS|React Native|Flutter|Xamarin|Ionic|Cordova|PhoneGap)\b',
            # Data Science
            r'\b(?:Pandas|NumPy|Scikit-learn|TensorFlow|PyTorch|Keras|Spark|Hadoop|Hive|Pig|Kafka|Airflow|Jupyter|Tableau|Power BI|Excel)\b',
            # Operating Systems
            r'\b(?:Linux|Windows|macOS|Unix|Ubuntu|CentOS|RedHat|Debian|Fedora)\b',
            # Methodologies
            r'\b(?:Agile|Scrum|Kanban|Waterfall|Lean|Six Sigma|ITIL|PMP|PRINCE2)\b',
            # Other technologies
            r'\b(?:API|REST|SOAP|GraphQL|JSON|XML|YAML|HTML|CSS|Nginx|Apache|Tomcat|IIS)\b'
        ]
        
        for pattern in tech_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            skills.update(matches)
        
        # 6. Clean and deduplicate skills
        cleaned_skills = []
        seen_skills = set()  # Track normalized skills to avoid duplicates
        
        for skill in skills:
            skill = skill.strip()
            if skill and len(skill) > 1:
                # Filter out date patterns (e.g., "Oct'2020 – September 2022", "From Feb 2006 To April 2019")
                if re.search(r'\d{4}', skill) and (re.search(r'(?:from|to|till|until|since)', skill, re.IGNORECASE) or 
                                                    re.search(r'[–-]\s*\d{4}', skill) or
                                                    re.search(r'\(.*\d{4}.*\)', skill)):
                    continue
                
                # Filter out location patterns with dates (e.g., "(OPTUM Technology Location: Eden Prairie Mn Usa (From April 2019 To June 2020))")
                if re.search(r'\(.*location.*\d{4}', skill, re.IGNORECASE):
                    continue
                
                # Filter out long sentences that are clearly not skills (more than 5 words typically indicates a sentence)
                words = skill.split()
                if len(words) > 5:
                    continue
                
                # Filter out text that looks like job descriptions or responsibilities
                if any(word in skill.lower() for word in ['responsibilities', 'worked', 'developed', 'created', 'implemented', 'leading', 'efforts']):
                    continue
                
                # Filter out "Environment:" patterns (e.g., "Environment: Apex Language", "Environment: Saleforce.Com Platform")
                if re.match(r'^Environment:\s*', skill, re.IGNORECASE):
                    continue
                if re.search(r'\bEnvironment:\s*[A-Z]', skill, re.IGNORECASE):
                    continue
                
                # Remove common prefixes/suffixes
                skill = re.sub(r'^(proficient|experienced|skilled|expert|knowledgeable|familiar)\s+(?:in|with|at)?\s*', '', skill, flags=re.IGNORECASE)
                skill = re.sub(r'\s+(proficient|experienced|skilled|expert|knowledgeable|familiar)$', '', skill, flags=re.IGNORECASE)
                
                # Normalize the skill for comparison (lowercase, remove extra spaces, handle common variations)
                normalized_skill = re.sub(r'\s+', ' ', skill.lower().strip())
                
                # Handle common variations
                normalized_skill = re.sub(r'\bnode\.js\b', 'nodejs', normalized_skill)
                normalized_skill = re.sub(r'\bci/cd\b', 'cicd', normalized_skill)
                normalized_skill = re.sub(r'\basp\.net\b', 'aspnet', normalized_skill)
                normalized_skill = re.sub(r'\bc\+\+\b', 'cpp', normalized_skill)
                normalized_skill = re.sub(r'\bc#\b', 'csharp', normalized_skill)
                
                # Only add if we haven't seen this normalized skill before
                if normalized_skill not in seen_skills and len(normalized_skill) > 1:
                    seen_skills.add(normalized_skill)
                    # Use the original case from the first occurrence
                    cleaned_skills.append(skill)
        
        # Sort the final list
        return sorted(cleaned_skills)

    # -------------------------------------------------------------------------
    # Experience extraction
    # -------------------------------------------------------------------------
    def extract_experience(self, experience_text: str) -> List[JobExperience]:
        """Extract work experience with improved parsing."""
        jobs = []
        lines = experience_text.split('\n')
        current_job = JobExperience()
        current_responsibilities = []

        def _parse_date_str(date_str: str) -> Optional[datetime]:
            """Best-effort parsing of a date string into a datetime (month-level granularity)."""
            if not date_str:
                return None
            s = date_str.strip()
            s_lower = s.lower()
            # Handle present/current variations (with and without spaces)
            if s_lower in {"present", "current", "now", "till date", "till now", "tilldate", "tillnow"}:
                return datetime.today()
            
            # Remove apostrophes from dates like "Oct'2020" -> "Oct 2020"
            s = re.sub(r"['']", " ", s)
            
            # Handle dates without spaces like "Jan2023" -> "Jan 2023"
            # Pattern: Month abbreviation followed directly by 4 digits
            s = re.sub(r'([A-Za-z]{3,9})(\d{4})', r'\1 \2', s)
            
            # Handle 2-digit years: "May22" or "JAN 22" -> "May 2022" or "JAN 2022"
            # Pattern: Month abbreviation followed by 2 digits (not 4)
            def convert_2digit_year(match):
                month = match.group(1)
                year_2digit = int(match.group(2))
                # Convert 2-digit year to 4-digit (00-30 -> 2000-2030, 31-99 -> 1931-1999)
                if year_2digit <= 30:
                    year_4digit = 2000 + year_2digit
                else:
                    year_4digit = 1900 + year_2digit
                return f"{month} {year_4digit}"
            
            s = re.sub(r'([A-Za-z]{3,9})\s*(\d{2})(?!\d)', convert_2digit_year, s, flags=re.IGNORECASE)
            
            s = re.sub(r'\s+', ' ', s).strip()

            # Try a series of common patterns
            candidates = [
                "%b %Y",      # Jan 2020, Oct 2020
                "%B %Y",      # January 2020
                "%m/%Y",      # 01/2020
                "%m-%Y",      # 01-2020
                "%Y/%m/%d",   # 2020/01/15
                "%Y-%m-%d",   # 2020-01-15
                "%m/%d/%Y",   # 01/15/2020
                "%d/%m/%Y",   # 15/01/2020
                "%Y",         # 2020 (assume Jan 1st)
            ]
            for fmt in candidates:
                try:
                    dt = datetime.strptime(s, fmt)
                    # Normalize to first of month to keep month-level comparisons simple
                    return datetime(dt.year, dt.month, 1)
                except Exception:
                    continue
            return None

        def _format_duration(start: Optional[datetime], end: Optional[datetime]) -> Optional[str]:
            """Return a human-readable duration between two datetimes."""
            if not start or not end:
                return None
            if end < start:
                return None
            total_months = (end.year - start.year) * 12 + (end.month - start.month)
            if total_months <= 0:
                return None
            years, months = divmod(total_months, 12)
            parts = []
            if years:
                parts.append(f"{years} year{'s' if years != 1 else ''}")
            if months:
                parts.append(f"{months} month{'s' if months != 1 else ''}")
            return " ".join(parts) if parts else None

        def _finalize_job(job: JobExperience):
            """Compute derived fields like duration before appending."""
            if job.start_date:
                start_dt = _parse_date_str(job.start_date)
            else:
                start_dt = None
            if job.end_date:
                end_dt = _parse_date_str(job.end_date)
            else:
                end_dt = None
            duration = _format_duration(start_dt, end_dt)
            if duration:
                job.duration = duration
            # If employment type not found in header/first line, search full description (responsibilities)
            if not job.employment_type and job.responsibilities:
                full_desc = "\n".join(job.responsibilities)
                job.employment_type = self._extract_employment_type(full_desc)
            # Don't extract department from full description (reduces false positives like "Design" from "Designing")
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if not line_stripped:
                continue
            
            # Check for date patterns
            date_match = DATE_RANGE_REGEX.search(line_stripped)
            
            if date_match:
                # Save previous job if exists
                if current_job.title or current_job.company:
                    current_job.responsibilities = current_responsibilities
                    _finalize_job(current_job)
                    jobs.append(current_job)
                
                # Start new job
                current_job = JobExperience()
                current_responsibilities = []
                
                # Extract dates
                date_text = date_match.group(0)
                dates = re.findall(r'(?i)(?:' + '|'.join(DATE_PATTERNS) + r'|present|current|till\s+date|till\s+now|tilldate|tillnow)', date_text)
                if len(dates) >= 1:
                    current_job.start_date = dates[0]
                if len(dates) >= 2:
                    current_job.end_date = dates[1]
                elif len(dates) == 1:
                    # Check if date_text contains "Current", "Present", etc. as end date
                    date_text_lower = date_text.lower()
                    if any(term in date_text_lower for term in ['current', 'present', 'now', 'till date', 'till now', 'tilldate', 'tillnow']):
                        # Find which term was used
                        for term in ['current', 'present', 'now', 'till date', 'till now', 'tilldate', 'tillnow']:
                            if term in date_text_lower:
                                current_job.end_date = term.title() if term not in ['till date', 'till now'] else term.title()
                                break
                
                # Pipe-separated table format: "Project Name | Client Name | Duration | Domain | Role" or "Project | Client | Duration | Role"
                if ' | ' in line_stripped:
                    parts = [p.strip() for p in line_stripped.split('|')]
                    # Find the segment that contains the date range we matched
                    date_seg_idx = None
                    for idx, seg in enumerate(parts):
                        if date_match.group(0) in seg or (current_job.start_date and current_job.start_date in seg):
                            date_seg_idx = idx
                            break
                    # Expected: 5 columns = Project | Client | Duration | Domain | Role
                    if date_seg_idx is not None and len(parts) >= 5 and date_seg_idx == 2:
                        current_job.title = parts[4].strip() if len(parts) > 4 else current_job.title  # Role
                        client = parts[1].strip() if len(parts) > 1 else ''
                        project = parts[0].strip() if len(parts) > 0 else ''
                        current_job.company = client if client and len(client) < 80 else (project if project and len(project) < 80 else current_job.company)
                    # 4 columns = Project | Client | Duration | Role
                    elif date_seg_idx is not None and len(parts) == 4 and date_seg_idx == 2:
                        current_job.title = parts[3].strip() if len(parts) > 3 else current_job.title
                        client = parts[1].strip() if len(parts) > 1 else ''
                        project = parts[0].strip() if len(parts) > 0 else ''
                        current_job.company = client if client and len(client) < 80 else (project if project and len(project) < 80 else current_job.company)
                    elif date_seg_idx is not None and len(parts) >= 3:
                        # 3+ columns, duration at some index - use last as role, second as company
                        current_job.title = parts[-1].strip() if parts and not current_job.title else current_job.title
                        if len(parts) >= 2 and not current_job.company:
                            current_job.company = parts[1].strip() if len(parts[1].strip()) < 80 else parts[0].strip()
                
                # Strategy 1: Extract title/company from the same line as date (skip if we already set from pipe format)
                date_start_pos = date_match.start()
                parsed_pipe = current_job.title and current_job.company and ' | ' in line_stripped
                if date_start_pos > 0 and not parsed_pipe:
                    # Extract text before the date
                    before_date = line_stripped[:date_start_pos].strip()
                    # Remove trailing comma if present
                    before_date = re.sub(r',\s*$', '', before_date).strip()
                    
                    # Check if before_date looks like a job title (contains job title keywords or "Role:")
                    # Handle "Role: Title" pattern
                    role_title_match = re.search(r'(?:Role|Position)[:\-]\s*(.+?)$', before_date, re.IGNORECASE)
                    if role_title_match:
                        title_candidate = role_title_match.group(1).strip()
                        if not current_job.title and len(title_candidate) < 100:
                            current_job.title = title_candidate
                    elif self._contains_job_title_keyword(before_date):
                        if not current_job.title and len(before_date) < 100:
                            # Remove "Role:" prefix if present
                            title_candidate = re.sub(r'^(?:Role|Position)[:\-]\s*', '', before_date, flags=re.IGNORECASE).strip()
                            current_job.title = title_candidate
                    # Check if it's a "Client: Company" pattern
                    elif re.search(r'^Client[:\-]\s*', before_date, re.IGNORECASE):
                        client_pattern = r'Client[:\-]\s*(.+?)(?:\s*[,\-]\s*[A-Z]{2})?\s*$'
                        client_match = re.search(client_pattern, before_date, re.IGNORECASE)
                        if client_match:
                            company_candidate = client_match.group(1).strip()
                            if company_candidate and len(company_candidate) < 100:
                                current_job.company = company_candidate
                    # Otherwise, check if it could be a company (but prioritize title extraction from next line)
                    elif not current_job.company and before_date and not self._looks_like_environment_or_tech_stack(before_date):
                        # Pattern: "Company, Location." or "Company, Location" (location is optional)
                        # Remove trailing period and location (City, ST or just ST)
                        company_candidate = before_date.rstrip('.')
                        # Remove location pattern: ", ST" or ", City, ST" or ", Country." (allow optional spaces after comma)
                        company_candidate = re.sub(r',\s*[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*,\s*[A-Z]{2}\s*$', '', company_candidate).strip()
                        company_candidate = re.sub(r',\s*[A-Z]{2}\s*$', '', company_candidate).strip()
                        company_candidate = re.sub(r',\s*[A-Za-z]{2,20}\.?\s*$', '', company_candidate).strip()  # ", India." or ", Country"
                        
                        # Only set as company if it doesn't look like a title or tech stack
                        if company_candidate and len(company_candidate) < 100 and not self._looks_like_environment_or_tech_stack(company_candidate):
                            words = company_candidate.split()
                            # Company: has company indicators, or (proper format and (2+ words or single long word) and not title)
                            if (re.search(r'\b(?:Inc|LLC|Corp|Co|Ltd|Company|Corporation|Bank|Services|Solutions|Technologies|Systems|Financial)\b', company_candidate, re.IGNORECASE) or
                                (re.match(r'^[A-Z][A-Za-z0-9\s&.,\-]+$', company_candidate) and
                                 (len(words) >= 2 or (len(words) == 1 and len(company_candidate) >= 12)) and
                                 not self._contains_job_title_keyword(company_candidate))):
                                current_job.company = company_candidate
                                
                                # Also extract location if present in before_date
                                location_match = re.search(r',\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*,\s*[A-Z]{2}|[A-Z]{2})\s*\.?$', before_date)
                                if location_match:
                                    current_job.location = location_match.group(1).strip()
                
                # Strategy 1.5: Look for title/company on the next line after date
                # Priority: Title first (common pattern: "Company, Location. Date\nTitle")
                if i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    if next_line and len(next_line) < 100 and not self._looks_like_environment_or_tech_stack(next_line):
                        # Skip if it's clearly a description or responsibility
                        if not (re.match(r'^(?:Description|Responsibilities?|Environment)[:\-]', next_line, re.IGNORECASE) or
                                self._compiled_patterns['bullet_point'].match(next_line)):
                            # Check for title keywords FIRST (titles often appear on next line)
                            if not current_job.title and self._contains_job_title_keyword(next_line):
                                current_job.title = next_line
                            # Then check for company indicators
                            elif not current_job.company and re.search(r'\b(?:Inc|LLC|Corp|Co|Ltd|Company|Corporation|Bank|Services|Solutions|Technologies|Systems)\b', next_line, re.IGNORECASE):
                                current_job.company = next_line
                            # Or if it looks like a company name (but not if it has title keywords)
                            elif not current_job.company and not current_job.title:
                                if (re.match(r'^[A-Z][A-Za-z0-9\s&.,\-]+$', next_line) and 
                                    len(next_line.split()) <= 5 and  # Company names are usually short
                                    not self._contains_job_title_keyword(next_line)):
                                    # Remove location if present
                                    company_candidate = re.sub(r',\s*[A-Z]{2}\s*$', '', next_line).strip()
                                    if company_candidate:
                                        current_job.company = company_candidate
                
                # Strategy 2: Look for company in previous lines (before date line)
                # Priority: Check line immediately before date (most common pattern: "Company, Location\nRole: Title Date")
                for j in range(max(0, i-3), i):
                    prev_line = lines[j].strip()
                    if not prev_line or self._looks_like_environment_or_tech_stack(prev_line):
                        continue
                    
                    # Skip if it's clearly a responsibility line (starts with bullet or common verbs)
                    if self._compiled_patterns['bullet_point'].match(prev_line) or self._compiled_patterns['responsibility_start'].match(prev_line):
                        continue
                    
                    # Skip if it contains "Role:" (this is a title line, not company)
                    if re.search(r'^Role[:\-]', prev_line, re.IGNORECASE):
                        continue
                    
                    # Pattern: "Company, City, ST" or "Company, ST" - extract company and location
                    location_pattern = r',\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*,\s*[A-Z]{2}|[A-Z]{2})\s*$'
                    location_match = re.search(location_pattern, prev_line)
                    if location_match:
                        # Extract company (everything before the location)
                        company_candidate = prev_line[:location_match.start()].strip()
                        # Extract location
                        location_text = location_match.group(1).strip()
                        if company_candidate and len(company_candidate) < 100:
                            # Validate it's a reasonable company name
                            if (re.match(r'^[A-Z][A-Za-z0-9\s&.,\-]+$', company_candidate) and 
                                len(company_candidate.split()) <= 5 and  # Company names are usually short
                                not any(keyword in company_candidate.lower() for keyword in JOB_TITLE_KEYWORDS)):
                                current_job.company = company_candidate
                                current_job.location = location_text
                                break
                    
                    # Pattern: "Client: Company Name" or "Client- Company Name"
                    client_pattern = r'Client[:\-]\s*(.+?)(?:\s*[,\-]|$)'
                    client_match = re.search(client_pattern, prev_line, re.IGNORECASE)
                    if client_match:
                        company_candidate = client_match.group(1).strip()
                        # Remove location if present (e.g., "Company, Location")
                        company_candidate = re.sub(r',\s*[A-Z]{2}\s*$', '', company_candidate).strip()
                        if company_candidate and len(company_candidate) < 100:
                            current_job.company = company_candidate
                            break
                    
                    # Check for company indicators (Inc, LLC, Corp, etc.)
                    if re.search(r'\b(?:Inc|LLC|Corp|Co|Ltd|Company|Corporation|Bank|Services|Solutions|Technologies|Systems)\b', prev_line, re.IGNORECASE):
                        if not current_job.company and len(prev_line) < 100:
                            # Remove "Client:" prefix if present
                            company_candidate = re.sub(r'^Client[:\-]\s*', '', prev_line, flags=re.IGNORECASE).strip()
                            # Remove location if present
                            company_candidate = re.sub(r',\s*[A-Z]{2}\s*$', '', company_candidate).strip()
                            if company_candidate:
                                current_job.company = company_candidate
                                break
                    
                    # Check if line contains job title keywords (this is likely a title, not company)
                    if self._contains_job_title_keyword(prev_line):
                        if not current_job.title:
                            current_job.title = prev_line
                    # Otherwise, if it looks like a company name (proper capitalization, no bullets, reasonable length)
                    elif not current_job.company and len(prev_line) < 100 and not self._compiled_patterns['bullet_point'].match(prev_line):
                        # Check if it's a reasonable company name format (2-5 words, proper capitalization)
                        words = prev_line.split()
                        if (2 <= len(words) <= 5 and 
                            re.match(r'^[A-Z][A-Za-z0-9\s&.,\-]+$', prev_line) and
                            not any(keyword in prev_line.lower() for keyword in JOB_TITLE_KEYWORDS)):
                            company_candidate = re.sub(r'^Client[:\-]\s*', '', prev_line, flags=re.IGNORECASE).strip()
                            company_candidate = re.sub(r',\s*[A-Z]{2}\s*$', '', company_candidate).strip()
                            if company_candidate:
                                current_job.company = company_candidate
                                break

                # Look for job title in lines after date (common pattern: "Role: Title" or just "Title")
                # Skip the immediate next line if we already extracted company from it
                start_idx = i + 2 if (i + 1 < len(lines) and lines[i + 1].strip() == current_job.company) else i + 1
                if not current_job.title:
                    for k in range(start_idx, min(len(lines), i + 4)):
                        title_line = lines[k].strip()
                        if not title_line or self._looks_like_environment_or_tech_stack(title_line):
                            continue
                        
                        # Skip if it's a responsibility line
                        if self._compiled_patterns['bullet_point'].match(title_line) or self._compiled_patterns['description_start'].match(title_line):
                            break
                        
                        # Skip if it's the company we already extracted
                        if title_line == current_job.company:
                            continue
                        
                        # Pattern: "Role: Job Title" or "Role- Job Title"
                        role_pattern = r'(?:Role|Position)[:\-]\s*(.+?)$'
                        role_match = re.search(role_pattern, title_line, re.IGNORECASE)
                        if role_match:
                            title_candidate = role_match.group(1).strip()
                            if title_candidate and len(title_candidate) < 100:
                                current_job.title = title_candidate
                                break
                        
                        # Check if line contains job title keywords
                        if self._contains_job_title_keyword(title_line):
                            if len(title_line) < 100 and not re.match(r'^(?:Description|Responsibilities?|Environment)', title_line, re.IGNORECASE):
                                current_job.title = title_line
                                break

                # Look for a location line in the few lines following the date line
                # Also check if location contains company info (e.g., "Client: Company, VA")
                    for k in range(i + 1, min(len(lines), i + 5)):
                        loc_line = lines[k].strip()
                        if not loc_line:
                            continue
                    
                    # Pattern: "Client: Company Name, Location"
                    client_location_pattern = r'Client[:\-]\s*(.+?),\s*([A-Z]{2})'
                    client_loc_match = re.search(client_location_pattern, loc_line, re.IGNORECASE)
                    if client_loc_match:
                        # Extract company from location field if not already set
                        if not current_job.company:
                            company_from_loc = client_loc_match.group(1).strip()
                            if company_from_loc and len(company_from_loc) < 100:
                                current_job.company = company_from_loc
                        # Set location
                        location_part = f"{client_loc_match.group(1)}, {client_loc_match.group(2)}"
                        current_job.location = location_part
                        break
                    
                    # Simple location pattern: City, ST
                        if re.search(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*,\s*[A-Z]{2}\b', loc_line):
                            current_job.location = loc_line
                            break

                # Extract employment type and department from header context (date line + next few lines)
                header_context = "\n".join(
                    [line_stripped]
                    + [lines[j].strip() for j in range(i + 1, min(i + 5, len(lines))) if j < len(lines) and lines[j].strip()]
                )
                if not current_job.employment_type:
                    current_job.employment_type = self._extract_employment_type(header_context)
                if not current_job.department:
                    current_job.department = self._extract_department(header_context)
            
            elif re.match(r'^(?:Roles?\s+and\s+Responsibilities|Technical\s+Skills|Professional\s+(?:Summary|Experience)|Education|Skills|Certifications|Projects?|Trainings?\s+&?\s*Certifications?)[\s:]', line_stripped, re.IGNORECASE):
                # Section header: finalize current job so we don't attach following content to last experience entry
                if current_job.title or current_job.company:
                    current_job.responsibilities = current_responsibilities
                    _finalize_job(current_job)
                    jobs.append(current_job)
                current_job = JobExperience()
                current_responsibilities = []
            elif current_job.title or current_job.company:
                # This is likely a responsibility or achievement
                if not current_job.employment_type or not current_job.department:
                    if not current_job.employment_type:
                        current_job.employment_type = self._extract_employment_type(line_stripped)
                    if not current_job.department:
                        current_job.department = self._extract_department(line_stripped)
                current_responsibilities.append(line_stripped)
        
        # Add the last job
        if current_job.title or current_job.company:
            current_job.responsibilities = current_responsibilities
            _finalize_job(current_job)
            jobs.append(current_job)
        
        # Deduplicate jobs - remove entries with same company and dates
        seen_jobs = set()
        deduplicated_jobs = []
        for job in jobs:
            # Create a unique key based on company and dates
            job_key = (
                (job.company or "").lower().strip(),
                (job.start_date or "").lower().strip(),
                (job.end_date or "").lower().strip()
            )
            # Skip if we've seen this exact combination before
            if job_key in seen_jobs and job_key[0]:  # Only deduplicate if company exists
                continue
            # Also skip if title and company are swapped (title contains company name and company contains title)
            if job.title and job.company:
                title_lower = job.title.lower()
                company_lower = job.company.lower()
                # Check if they might be swapped
                if company_lower in title_lower or title_lower in company_lower:
                    # If company looks more like a title (has job title keywords), swap them
                    if self._contains_job_title_keyword(company_lower) and not self._contains_job_title_keyword(title_lower):
                        job.title, job.company = job.company, job.title
                    # If title looks more like a company (has company indicators), swap them
                    elif self._compiled_patterns['company_indicators'].search(title_lower) and not self._compiled_patterns['company_indicators'].search(company_lower):
                        job.title, job.company = job.company, job.title
            
            # Clean up title: remove "Role:" prefix if present (using cached pattern)
            if job.title:
                job.title = self._compiled_patterns['role_prefix'].sub('', job.title).strip()
            
            # Don't keep Environment/tech-stack lines as company or title
            if job.company and self._looks_like_environment_or_tech_stack(job.company):
                job.company = ""
            if job.title and self._looks_like_environment_or_tech_stack(job.title):
                job.title = ""
            
            seen_jobs.add(job_key)
            deduplicated_jobs.append(job)
        
        # Drop certification rows, table headers, and entries with neither company nor title
        filtered_jobs = [
            j for j in deduplicated_jobs
            if not self._is_certification_or_table_entry(j)
            and (j.company or j.title)  # keep if at least one of company or title is set
        ]
        return filtered_jobs

    # -------------------------------------------------------------------------
    # Education extraction
    # -------------------------------------------------------------------------
    def extract_education(self, education_text: str) -> List[Education]:
        """Extract education information."""
        education_list = []
        lines = education_text.split('\n')
        
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped or len(line_stripped) < 10:
                continue
            
            # Skip lines that are clearly not education (Environment, Tools, Technologies, etc.)
            if re.search(r'^(?:Environment|Tools|Technologies|Skills|Technical|Programming|Languages|Frameworks|Databases|Platforms|EXPERIENCE|TECHNICAL\s+SKILLS)[:\-]', line_stripped, re.IGNORECASE):
                continue
            
            # Skip lines that are section headers (all caps, short lines)
            if line_stripped.isupper() and len(line_stripped.split()) <= 3:
                continue
            
            # Skip lines that start with (cid:0) and don't contain education keywords
            if re.match(r'^\(cid:\d+\)', line_stripped) and not re.search(r'\b(?:University|College|Institute|School|Bachelor|Master|PhD|Degree)\b', line_stripped, re.IGNORECASE):
                continue
            
            # Skip lines that are clearly responsibility/experience lines
            if re.match(r'^(?:\(cid:\d+\)\s*)?(?:Worked|Developed|Created|Implemented|Designed|Managed|Responsible|Involved|Experience|Skills|Tools|Environment)', line_stripped, re.IGNORECASE):
                continue
            
            # Skip lines that are clearly certifications (avoid "GCP Professional Data Engineer certified" as education)
            if re.search(r'\b(?:certified|certification|certificate)\b', line_stripped, re.IGNORECASE):
                continue
            # Skip technical/skills lines misclassified as education (e.g. "Operating System\tWindows, Linux")
            if re.search(r'^(?:Operating\s+System|RDBMS|Languages|Web\s+Technologies|Development\s+Tools)[\s:\t]', line_stripped, re.IGNORECASE):
                continue
            
            # Look for degree patterns (including plural and possessive; allow Unicode apostrophe in Master's/Bachelor's)
            degree_patterns = [
                r'\b(?:Bachelor|Bachelor[\'\u2019]?s|Master|Master[\'\u2019]?s|Masters|PhD|Ph\.?D\.?|Doctorate|Associate|Certificate|Diploma)\b',
                r'\b(?:B\.?S\.?|M\.?S\.?|B\.?A\.?|M\.?A\.?|Ph\.?D\.?)\b'
            ]
            
            if any(re.search(pattern, line_stripped, re.IGNORECASE) for pattern in degree_patterns):
                # Don't treat "Associate" in certification phrases as degree (e.g. "Solutions Architect – Associate")
                if re.search(r'\bAssociate\b', line_stripped, re.IGNORECASE):
                    if re.search(r'(?:certified|certification|certificate).*\bAssociate\b|(?:architect|engineer)\s*[–\-]\s*Associate\b', line_stripped, re.IGNORECASE):
                        continue
                
                education = Education()
                
                # Extract degree
                for pattern in degree_patterns:
                    match = re.search(pattern, line_stripped, re.IGNORECASE)
                    if match:
                        degree_text = match.group(0)
                        # Normalize degree text
                        if degree_text.lower() in ['masters', "master's"]:
                            education.degree = "Master"
                        elif degree_text.lower() in ["bachelor's", 'bachelors']:
                            education.degree = "Bachelor"
                        else:
                            education.degree = degree_text
                        break
                
                # Extract institution, field, and graduation date
                # Clean the line: remove bullet points and extra whitespace
                cleaned_line = re.sub(r'^[●•\-\*]\s*', '', line_stripped).strip()
                
                # Pattern 1a: "Master's in X, from Y – Year" (with comma and dash)
                field_institution_pattern_comma = r'(?:in|of)\s+([A-Z][A-Za-z\s]+?),\s+from\s+([A-Z][A-Za-z0-9\s]+?)\s*[–\-]\s*(\d{4})'
                field_match_comma = re.search(field_institution_pattern_comma, cleaned_line, re.IGNORECASE)
                if field_match_comma:
                    education.field = field_match_comma.group(1).strip()
                    institution_text = field_match_comma.group(2).strip()
                    institution_text = re.sub(r'[.,;]$', '', institution_text).strip()
                    education.institution = institution_text
                    education.graduation_date = field_match_comma.group(3).strip()
                else:
                    # Pattern 1b: "Bachelor in ECE from JNTUH in 2014" or "Master in CS from University in 2020"
                    # Handle both acronyms (all caps) and regular words
                    field_institution_pattern = r'(?:in|of)\s+([A-Z][A-Za-z\s]+?)\s+from\s+([A-Z][A-Za-z0-9\s]+?)(?:\s+in\s+(\d{4}))?(?:\s|$)'
                    field_match = re.search(field_institution_pattern, cleaned_line, re.IGNORECASE)
                    if field_match:
                        education.field = field_match.group(1).strip()
                        institution_text = field_match.group(2).strip()
                        # Clean up institution - remove trailing punctuation and ensure we get the full name
                        institution_text = re.sub(r'[.,;]$', '', institution_text).strip()
                        education.institution = institution_text
                        if field_match.group(3):
                            education.graduation_date = field_match.group(3).strip()
                    else:
                        # Pattern 2: "Masters: Computer Science" - extract what comes after the colon
                        colon_match = re.search(r'[:\-]\s*(.+)', cleaned_line)
                        if colon_match:
                            after_colon = colon_match.group(1).strip()
                            # Try to extract field and institution from "Computer Science" or "CS from University"
                            field_inst_match = re.search(r'^([A-Z][A-Za-z\s]+?)(?:\s+from\s+([A-Z][A-Za-z\s]+?))?(?:\s+in\s+(\d{4}))?$', after_colon, re.IGNORECASE)
                            if field_inst_match:
                                potential_field = field_inst_match.group(1).strip()
                                potential_inst = field_inst_match.group(2).strip() if field_inst_match.group(2) else None
                                potential_year = field_inst_match.group(3).strip() if field_inst_match.group(3) else None
                                
                                # Check if potential_field contains institution keywords
                                if re.search(r'\b(?:University|College|Institute|School)\b', potential_field, re.IGNORECASE):
                                    education.institution = potential_field
                                else:
                                    education.field = potential_field
                                    if potential_inst:
                                        education.institution = potential_inst
                                    else:
                                        # Look for institution in next line
                                        line_idx = lines.index(line)
                                        if line_idx + 1 < len(lines):
                                            next_line = lines[line_idx + 1].strip()
                                            if re.search(r'\b(?:University|College|Institute|School)\b', next_line, re.IGNORECASE):
                                                education.institution = next_line
                                            elif next_line:
                                                education.institution = next_line
                                
                                if potential_year:
                                    education.graduation_date = potential_year
                            else:
                                # If it contains University/College/Institute, use it as institution
                                if re.search(r'\b(?:University|College|Institute|School)\b', after_colon, re.IGNORECASE):
                                    education.institution = after_colon
                                else:
                                    # Otherwise, look for institution in next line
                                    line_idx = lines.index(line)
                                    if line_idx + 1 < len(lines):
                                        next_line = lines[line_idx + 1].strip()
                                        if re.search(r'\b(?:University|College|Institute|School)\b', next_line, re.IGNORECASE):
                                            education.institution = next_line
                                        else:
                                            education.institution = next_line if next_line else after_colon
                        elif re.search(r'\b(?:University|College|Institute|School)\b', cleaned_line, re.IGNORECASE):
                            education.institution = cleaned_line
                        else:
                            # Look for institution in next line
                            line_idx = lines.index(line)
                            if line_idx + 1 < len(lines):
                                next_line = lines[line_idx + 1].strip()
                                if next_line and len(next_line) > 3:
                                    education.institution = next_line
                            else:
                                education.institution = cleaned_line
                
                # Extract graduation date if not already extracted (look for year in the line)
                if not education.graduation_date:
                    year_match = re.search(r'\b(19|20)\d{2}\b', cleaned_line)
                    if year_match:
                        education.graduation_date = year_match.group(0)
                
                # Extract GPA if present (various formats: "GPA: 3.8", "3.8/4.0", "CGPA: 8.5", etc.)
                if not education.gpa:
                    # Comprehensive GPA patterns
                    gpa_patterns = [
                        r'(?:GPA|CGPA)[:\-]?\s*(\d+\.?\d*)',  # "GPA: 3.8" or "GPA 3.8" or "CGPA: 8.5"
                        r'(?:GPA|CGPA)[:\-]?\s*(\d+\.?\d*)\s*/\s*\d+\.?\d*',  # "GPA: 3.8/4.0"
                        r'(?:GPA|CGPA)[:\-]?\s*(\d+\.?\d*)\s*(?:out\s+of|/)\s*\d+\.?\d*',  # "GPA 3.8 out of 4.0"
                        r'(\d+\.?\d*)\s*/\s*\d+\.?\d*',  # "3.8/4.0" or "8.5/10" (standalone)
                        r'(\d+\.?\d*)\s*(?:GPA|CGPA)',  # "3.8 GPA"
                        r'(\d+\.?\d*)\s*(?:GPA|CGPA)\s*\([^)]*\)',  # "3.8 GPA (out of 4.0)"
                    ]
                    
                    # Check current line first
                    for pattern in gpa_patterns:
                        gpa_match = re.search(pattern, cleaned_line, re.IGNORECASE)
                        if gpa_match:
                            gpa_value = gpa_match.group(1).strip()
                            # Validate GPA is reasonable (typically 0.0-10.0 or 0.0-4.0)
                            try:
                                gpa_float = float(gpa_value)
                                if 0.0 <= gpa_float <= 10.0:
                                    education.gpa = gpa_value
                                    break
                            except ValueError:
                                continue
                    
                    # If not found in current line, check surrounding lines (next 2 lines)
                    if not education.gpa:
                        line_idx = lines.index(line)
                        # Check next 1-2 lines (GPA is usually close to education info)
                        for offset in [1, 2]:
                            if line_idx + offset < len(lines):
                                check_line = lines[line_idx + offset].strip()
                                # Skip if line is too long (likely not GPA) or contains section headers
                                if (check_line and len(check_line) < 50 and 
                                    not re.search(r'^(?:EDUCATION|EXPERIENCE|SKILLS|CERTIFICATIONS)', check_line, re.IGNORECASE)):
                                    for pattern in gpa_patterns:
                                        gpa_match = re.search(pattern, check_line, re.IGNORECASE)
                                        if gpa_match:
                                            gpa_value = gpa_match.group(1).strip()
                                            try:
                                                gpa_float = float(gpa_value)
                                                if 0.0 <= gpa_float <= 10.0:
                                                    education.gpa = gpa_value
                                                    break
                                            except ValueError:
                                                continue
                                    if education.gpa:
                                        break
                
                # Validate education entry before adding - must have valid institution
                if education.institution:
                    institution_lower = education.institution.lower()
                    # Must contain education keywords OR be a reasonable institution name (2+ words, proper capitalization)
                    has_edu_keyword = re.search(r'\b(?:University|College|Institute|School|Univ|Tech|State|National)\b', institution_lower)
                    is_reasonable_name = (
                        len(education.institution.split()) >= 2 and
                        re.match(r'^[A-Z][A-Za-z\s&.,\-]+$', education.institution) and
                        not re.search(r'\b(?:Environment|Tools|Technologies|Skills|Technical|Programming|Languages|Frameworks|Databases|Platforms|EXPERIENCE|TECHNICAL\s+SKILLS|Worked|Developed|Created|Implemented)\b', institution_lower)
                    )
                    if has_edu_keyword or is_reasonable_name:
                        education_list.append(education)
        
        return education_list

    # -------------------------------------------------------------------------
    # Confidence score & projects
    # -------------------------------------------------------------------------
    def calculate_confidence_score(self, parsed_resume: ParsedResume) -> float:
        """Calculate confidence score based on extracted information."""
        score = 0.0
        max_score = 10.0
        
        # Contact information (3 points)
        if parsed_resume.contact.name:
            score += 1.0
        if parsed_resume.contact.email:
            score += 1.0
        if parsed_resume.contact.phone:
            score += 1.0
        
        # Professional summary (1 point)
        if parsed_resume.professionalSummary:
            score += 1.0
        
        # Experience (3 points)
        if parsed_resume.experience:
            score += min(3.0, len(parsed_resume.experience) * 0.5)
        
        # Education (1 point)
        if parsed_resume.education:
            score += 1.0
        
        # Skills (1 point)
        if parsed_resume.skills:
            score += 1.0
        
        # Certifications (1 point)
        if parsed_resume.certifications:
            score += 1.0
        
        return min(1.0, score / max_score)

    def extract_projects(self, text: str, sections: Optional[Dict[str, str]] = None) -> List[str]:
        """Extract project information from resume.
        
        Looks for project entries in the projects section or identifies project
        patterns throughout the document.
        """
        projects = []
        
        # Get projects section if available
        if sections is None:
            sections = self.find_sections(text)
        
        projects_section = sections.get("projects", "")
        
        if projects_section:
            # Parse project entries from the section
            lines = projects_section.split('\n')
            current_project = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    # Empty line might separate projects
                    if current_project:
                        project_text = ' '.join(current_project).strip()
                        if project_text and len(project_text) > 10:
                            projects.append(project_text)
                        current_project = []
                    continue
                
                # Check if this line starts a new project
                # Common patterns: "Project Name:", "Project: Name", "ProjectName -", bold/caps project names
                is_new_project = False
                
                # Pattern 1: Line starts with "Project" keyword
                if re.match(r'^(?:Project|Client)[:\s-]+', line, re.IGNORECASE):
                    is_new_project = True
                
                # Pattern 2: Line is mostly uppercase or title case (likely a project name)
                elif (line.isupper() or 
                      (line[0].isupper() and ':' in line[:30]) or
                      re.match(r'^[A-Z][A-Za-z0-9\s&-]+(?::|$)', line)):
                    # Check if it looks like a project name (not a description)
                    if len(line) < 100 and not line.lower().startswith(('responsible', 'worked', 'developed', 'created', 'implemented', 'designed', 'built', 'managed')):
                        is_new_project = True
                
                # Pattern 3: Bullet point with capitalized text (new project entry)
                elif re.match(r'^[•\-\*]\s*[A-Z]', line):
                    # Could be a new project or a description point
                    # If it's short and doesn't start with action verbs, likely a project name
                    clean_line = re.sub(r'^[•\-\*]\s*', '', line)
                    if len(clean_line) < 60 and not any(clean_line.lower().startswith(v) for v in 
                        ['responsible', 'worked', 'developed', 'created', 'implemented', 'designed', 'built', 'managed', 'analyzed', 'prepared']):
                        is_new_project = True
                
                if is_new_project:
                    # Save previous project
                    if current_project:
                        project_text = ' '.join(current_project).strip()
                        if project_text and len(project_text) > 10:
                            projects.append(project_text)
                    current_project = [line]
                else:
                    # Continue current project description
                    current_project.append(line)
            
            # Don't forget the last project
            if current_project:
                project_text = ' '.join(current_project).strip()
                if project_text and len(project_text) > 10:
                    projects.append(project_text)
        
        # If no projects found in section, try to find project patterns in text
        if not projects:
            # Look for "Project:" or "Project Name:" patterns
            project_patterns = [
                r'(?:Project|Client)[:\s-]+([A-Z][A-Za-z0-9\s&-]+?)(?:\n|:)',
                r'\b([A-Z][A-Za-z0-9\s&-]+)\s*[-–:]\s*(?:Worked|Responsible|Developed|Built|Created)',
            ]
            
            for pattern in project_patterns:
                matches = re.findall(pattern, text)
                for match in matches:
                    if isinstance(match, str):
                        project_name = match.strip()
                        if project_name and 3 < len(project_name) < 100:
                            projects.append(project_name)
        
        # Clean up and deduplicate
        cleaned_projects = []
        seen = set()
        for project in projects:
            # Clean up the project text
            project = re.sub(r'\s+', ' ', project).strip()
            # Remove trailing colons
            project = project.rstrip(':')
            
            # Skip if too short, duplicate, or looks like a section header
            project_lower = project.lower()
            if (len(project) < 10 or 
                project_lower in seen or
                project_lower in ['projects', 'key projects', 'recent projects', 'project experience']):
                continue
            
            seen.add(project_lower)
            cleaned_projects.append(project)
        
        return cleaned_projects

    # -------------------------------------------------------------------------
    # Main entry — parse_resume
    # -------------------------------------------------------------------------
    def parse_resume(self, file_path: str) -> ParsedResume:
        """Parse resume from file with comprehensive extraction."""
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Extract text based on file type
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            text = self.extract_text_from_pdf(file_path)
        elif ext == ".docx":
            text = self.extract_text_from_docx(file_path)
        else:
            raise ValueError("Unsupported file type. Only PDF and DOCX are supported.")
        
        if not text:
            logger.warning(f"No text extracted from {file_path}")
            return ParsedResume(file_path=file_path, contact=ContactInfo())
        text = self._normalize_raw_text(text)

        # Extract sections
        sections = self.find_sections(text)
        
        # Extract contact information
        contact = self.extract_contact_info(text)

        # Professional summary: prefer section content, then fallback regex on full text
        professional_summary = ""
        professional_summary_section = sections.get("professionalSummary", "")
        if professional_summary_section and len(professional_summary_section.strip()) > 20:
            # Clean up: remove section header if present
            professional_summary = professional_summary_section
            # Remove common header patterns
            professional_summary = re.sub(r'^(?:PROFESSIONAL\s+SUMMARY|SUMMARY|PROFILE|CAREER\s+SUMMARY|PROFESSIONAL\s+PROFILE|EXECUTIVE\s+SUMMARY|OBJECTIVE|CAREER\s+OBJECTIVE)[:\-]?\s*', '', professional_summary, flags=re.IGNORECASE)
            # Stop at next section headers (EDUCATION, EXPERIENCE, TECHNICAL SKILLS, etc.)
            # Only stop if the section header is on its own line (not in the middle of text)
            next_section_pattern = r'\n\s*(?:EDUCATION|EXPERIENCE|TECHNICAL\s+SKILLS|SKILLS|WORK\s+HISTORY|PROFESSIONAL\s+EXPERIENCE|WORKING\s+EXPERIENCE|CERTIFICATIONS|PROJECTS|AWARDS)\s*(?:\n|:)'
            next_section_match = re.search(next_section_pattern, professional_summary, re.IGNORECASE)
            if next_section_match:
                professional_summary = professional_summary[:next_section_match.start()].strip()
            professional_summary = professional_summary.strip()
        if not professional_summary:
            # Fallback: search entire text for professional summary
            summary_patterns = [
                r'(?:PROFESSIONAL\s+SUMMARY|SUMMARY|PROFILE|CAREER\s+SUMMARY|PROFESSIONAL\s+PROFILE|EXECUTIVE\s+SUMMARY|OBJECTIVE|CAREER\s+OBJECTIVE)[:\-]?\s*\n(.*?)(?=\n\s*(?:TECHNICAL\s+SKILLS|SKILLS|EXPERIENCE|EDUCATION|WORK\s+HISTORY|PROFESSIONAL\s+EXPERIENCE|CERTIFICATIONS)|$)',
            ]
            for pattern in summary_patterns:
                match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
                if match:
                    content = match.group(1).strip()
                    if len(content) > 20:
                        professional_summary = content
                        break
        # Use experience section if found, otherwise search entire text
        experience_section = sections.get("experience", "")
        if not experience_section or len(experience_section.strip()) < 50:
            # Fallback: search entire text for experience patterns
            experience = self.extract_experience(text)
        else:
            experience = self.extract_experience(experience_section)
            # If section had no/few jobs but full text has date ranges, section may be wrong or truncated
            # (e.g. jobs under "Projects/Technologies Experience" went to projects). Use full text.
            date_range_count = len(DATE_RANGE_REGEX.findall(text))
            if date_range_count > len(experience) and date_range_count >= 1:
                experience_full = self.extract_experience(text)
                if len(experience_full) > len(experience):
                    experience = experience_full
        # Use education section if found, otherwise search entire text
        education_section = sections.get("education", "")
        if not education_section or len(education_section.strip()) < 20:
            # Fallback: search entire text for education patterns
            education = self.extract_education(text)
        else:
            education = self.extract_education(education_section)
        # Pass precomputed sections into extract_skills to avoid redundant work
        skills = self.extract_skills(text, sections)
        # Extract certifications from multiple sources
        certifications = []
        
        # From certifications section
        cert_section = sections.get("certifications", "")
        if cert_section:
            for line in cert_section.split('\n'):
                line = line.strip()
                if line and any(word in line.lower() for word in 
                               ["certified", "certificate", "certification", "license", "credential"]):
                    certifications.append(line)
        
        # From entire text (look for certification patterns)
        cert_patterns = [
            r'\b(?:AWS|Azure|Google Cloud|Salesforce|TOSCA|PMP|ITIL|CISSP|CISA|CISM)\s+(?:Certified|Certification|Certificate)\b',
            r'\b(?:Certified|Certification|Certificate)\s+(?:in|for)?\s*(?:AWS|Azure|Google Cloud|Salesforce|TOSCA|PMP|ITIL|CISSP|CISA|CISM)\b',
            r'\b(?:AWS|Azure|Google Cloud|Salesforce|TOSCA|PMP|ITIL|CISSP|CISA|CISM)\s+\w+\s+(?:Certified|Certification|Certificate)\b'
        ]
        
        for pattern in cert_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            certifications.extend(matches)
        
        # Remove duplicates and clean up
        certifications = list(set([cert.strip() for cert in certifications if cert.strip()]))
        
        # Extract projects
        projects = self.extract_projects(text, sections)
        
        # Create parsed resume object
        parsed_resume = ParsedResume(
            file_path=file_path,
            contact=contact,
            professionalSummary=professional_summary,
            experience=experience,
            education=education,
            skills=skills,
            certifications=certifications,
            projects=projects,
            raw_text=text
        )
        
        # Calculate confidence score
        parsed_resume.confidence_score = self.calculate_confidence_score(parsed_resume)
        
        return parsed_resume

    # -------------------------------------------------------------------------
    # Serialization — duration, dates, to_dict / response shape
    # -------------------------------------------------------------------------
    def _parse_duration_to_months(self, duration_str: Optional[str]) -> Optional[int]:
        """Parse duration string like '5 years', '6 months', '2 years 3 months' to total months. Returns None if unparseable."""
        if not duration_str or not duration_str.strip():
            return None
        s = duration_str.strip().lower()
        total = 0
        # e.g. "2 years 3 months", "5 years", "6 months"
        years_m = re.search(r"(\d+)\s*year", s)
        if years_m:
            total += int(years_m.group(1)) * 12
        months_m = re.search(r"(\d+)\s*month", s)
        if months_m:
            total += int(months_m.group(1))
        return total if total else None

    def _normalize_date_to_iso(self, date_str: Optional[str]) -> Optional[str]:
        """Try to normalize date string to YYYY-MM-DD. Returns original string or None if empty."""
        if not date_str or not date_str.strip():
            return None
        s = date_str.strip()
        # Normalize apostrophes and collapse spaces
        s = re.sub(r"['\u2018\u2019]", " ", s)
        s = re.sub(r"\s+", " ", s).strip()
        # Expand month+2-digit year (e.g. "Oct 22", "May'22") to 4-digit year
        def two_digit_year(match):
            mo, yr = match.group(1), match.group(2)
            y = int(yr)
            year_4 = 2000 + y if y <= 50 else 1900 + y  # 00-50 -> 2000-2050, 51-99 -> 1951-1999
            return f"{mo} {year_4}"
        s = re.sub(r"([A-Za-z]{3,9})\s*['\s]*(\d{2})(?!\d)", two_digit_year, s, flags=re.IGNORECASE)
        for fmt in ("%Y-%m-%d", "%Y-%m", "%m/%d/%Y", "%m/%Y", "%d/%m/%Y", "%b %Y", "%B %Y", "%b %y", "%B %y", "%Y"):
            try:
                dt = datetime.strptime(s, fmt)
                return dt.strftime("%Y-%m-%d")
            except ValueError:
                continue
        return s

    def _experience_item_to_response(self, job: JobExperience) -> Dict[str, Any]:
        """Build one experience item with fixed keys: company.name, employmentType, title, department, startDate, endDate, isCurrent, durationInMonths, description."""
        end_raw = (job.end_date or "").strip().lower()
        is_current = not job.end_date or not job.end_date.strip() or end_raw in (
            "present", "current", "now", "till date", "till now", "tilldate", "tillnow", "to present"
        )
        start_date = self._normalize_date_to_iso(job.start_date) if job.start_date else None
        end_date = None if is_current else (self._normalize_date_to_iso(job.end_date) if job.end_date else None)
        duration_months = self._parse_duration_to_months(job.duration)
        # If duration string wasn't parseable, compute from start/end dates when possible
        if duration_months is None and start_date:
            try:
                start_dt = datetime.strptime(start_date, "%Y-%m-%d")
                end_dt = datetime.today() if is_current else (datetime.strptime(end_date, "%Y-%m-%d") if end_date else None)
                if end_dt and end_dt >= start_dt:
                    duration_months = (end_dt.year - start_dt.year) * 12 + (end_dt.month - start_dt.month)
                    if duration_months < 0:
                        duration_months = None
            except (ValueError, TypeError):
                pass
        description = "\n".join(job.responsibilities).strip() if job.responsibilities else ""
        # Don't use date-range-only or location-style text as title in response
        title_out = job.title or ""
        if self._is_date_range_only(title_out) or self._is_location_or_context_title(title_out):
            title_out = ""
        return {
            "company": {"name": job.company or ""},
            "employmentType": job.employment_type or "",
            "title": title_out,
            "department": job.department or "",
            "startDate": start_date,
            "endDate": end_date,
            "isCurrent": is_current,
            "durationInMonths": duration_months,
            "description": description,
        }

    def _contact_to_response(self, contact: ContactInfo) -> Dict[str, Any]:
        """Build contact object: firstName, middleName, lastName, fullName, email, phone, linkedinUrl, githubUrl, whatsapp, dateOfBirth, facebook, gender. Missing values are empty strings."""
        name = (contact.name or "").strip()
        parts = name.split() if name else []
        if len(parts) == 0:
            first_name, middle_name, last_name = "", "", ""
        elif len(parts) == 1:
            first_name, middle_name, last_name = parts[0], "", ""
        else:
            first_name = parts[0]
            last_name = parts[-1]
            middle_name = " ".join(parts[1:-1]) if len(parts) > 2 else ""
        return {
            "firstName": first_name,
            "middleName": middle_name,
            "lastName": last_name,
            "fullName": name,
            "email": contact.email or "",
            "phone": contact.phone or "",
            "linkedinUrl": contact.linkedin or "",
            "githubUrl": contact.github or "",
            "whatsapp": contact.whatsapp or "",
            "dateOfBirth": contact.date_of_birth or "",
            "facebook": contact.facebook or "",
            "gender": contact.gender or "",
        }

    def to_dict(self, parsed_resume: ParsedResume) -> Dict[str, Any]:
        """Convert ParsedResume to dictionary for JSON serialization."""
        result = asdict(parsed_resume)
        
        # Replace contact with fixed response shape (firstName, middleName, lastName, fullName, email, phone, linkedinUrl, githubUrl, whatsapp, dateOfBirth, facebook, gender)
        if result.get("contact"):
            result["contact"] = self._contact_to_response(parsed_resume.contact)
        
        # Replace experience with fixed response shape (company.name, employmentType, title, department, startDate, endDate, isCurrent, durationInMonths, description)
        exp_list = parsed_resume.experience or []
        result["experience"] = [self._experience_item_to_response(j) for j in exp_list]
        
        # Add experience count above the experience section
        result["experience_count"] = len(result["experience"])
        keys = list(result.keys())
        keys.remove("experience_count")
        idx = keys.index("experience") if "experience" in keys else len(keys)
        keys.insert(idx, "experience_count")
        result = {k: result[k] for k in keys}
        
        # Remove None location and gpa values from education entries
        if result.get("education"):
            for edu in result["education"]:
                if edu.get("location") is None:
                    edu.pop("location", None)
                if edu.get("gpa") is None:
                    edu.pop("gpa", None)
        
        return result

